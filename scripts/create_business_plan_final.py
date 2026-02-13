# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)

# 스크린샷 경로
screenshot_dir = r'G:\내 드라이브\developer\hanmed-cdss\docs\screenshots'

def add_bold_text(paragraph, text):
    """볼드 텍스트 추가"""
    run = paragraph.add_run(text)
    run.bold = True
    return run

def add_normal_text(paragraph, text):
    """일반 텍스트 추가"""
    run = paragraph.add_run(text)
    return run

def add_mixed_paragraph(doc, parts):
    """볼드와 일반 텍스트가 섞인 단락 생성
    parts: [('text', is_bold), ...]
    """
    p = doc.add_paragraph()
    for text, is_bold in parts:
        run = p.add_run(text)
        run.bold = is_bold
    return p

# Title
title = doc.add_heading('초기창업패키지 창업기업 사업계획서', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Notice
notice = doc.add_paragraph()
run = notice.add_run('※ 사업계획서는 목차(1페이지)를 제외하고 15페이지 이내로 작성')
run.font.size = Pt(9)

doc.add_paragraph()

# ========================================
# Section 1: 일반현황
# ========================================
doc.add_heading('□ 일반현황', level=1)

table1 = doc.add_table(rows=6, cols=4)
table1.style = 'Table Grid'

# Row 1
table1.cell(0, 0).text = '기업명'
cell = table1.cell(0, 1)
p = cell.paragraphs[0]
add_bold_text(p, '머프키치')

table1.cell(0, 2).text = '개업연월일'
table1.cell(0, 3).text = '22.01.31'

# Row 2
table1.cell(1, 0).text = '사업자 구분'
table1.cell(1, 1).text = '개인사업자'
table1.cell(1, 2).text = '대표자 유형'
table1.cell(1, 3).text = '단독'

# Row 3
table1.cell(2, 0).text = '사업자등록번호'
table1.cell(2, 1).text = '401-20-84647'
table1.cell(2, 2).text = '사업장 소재지'
table1.cell(2, 3).text = '경기도 의왕시 학의동'

# Row 4: 창업아이템명 - 핵심!
table1.cell(3, 0).text = '창업아이템명'
cell = table1.cell(3, 1)
cell.merge(table1.cell(3, 3))
p = cell.paragraphs[0]
add_bold_text(p, 'Claude AI + RAG 기술')
add_normal_text(p, ' 기반 ')
add_bold_text(p, '한의학 임상의사결정지원시스템(CDSS)')
add_normal_text(p, ' - ')
add_bold_text(p, '국내 최초 한의학 AI')
add_normal_text(p, ' 진료 지원 플랫폼')

# Row 5: 산출물
table1.cell(4, 0).text = '산출물\n(협약기간 내 목표)'
cell = table1.cell(4, 1)
cell.merge(table1.cell(4, 3))
p = cell.paragraphs[0]
add_bold_text(p, 'AI CDSS 웹 플랫폼(1개)')
add_normal_text(p, ', ')
add_bold_text(p, 'AI 추론 엔진(1개)')
add_normal_text(p, ', 환자용 모바일 앱(1개)')

# Row 6
table1.cell(5, 0).text = '지원 분야'
cell = table1.cell(5, 1)
cell.merge(table1.cell(5, 3))
p = cell.paragraphs[0]
add_bold_text(p, '☑ 지식서비스 (헬스케어 AI)')
add_normal_text(p, '   ☐ 제조')

doc.add_paragraph()

# 추가 테이블
table2 = doc.add_table(rows=3, cols=4)
table2.style = 'Table Grid'

table2.cell(0, 0).text = '전문기술분야'
cell = table2.cell(0, 1)
cell.merge(table2.cell(0, 3))
p = cell.paragraphs[0]
add_bold_text(p, '☑ 정보·통신 (인공지능, 빅데이터)')

table2.cell(1, 0).text = '총 사업비\n구성 계획'
cell = table2.cell(1, 1)
p = cell.paragraphs[0]
add_normal_text(p, '정부지원사업비\n')
add_bold_text(p, '100백만원')

cell = table2.cell(1, 2)
p = cell.paragraphs[0]
add_normal_text(p, '자기부담사업비\n현금: 14.29백만원\n현물: 28.57백만원')

cell = table2.cell(1, 3)
p = cell.paragraphs[0]
add_normal_text(p, '총 사업비\n')
add_bold_text(p, '142.86백만원')

table2.cell(2, 0).text = '지방우대 지역'
cell = table2.cell(2, 1)
cell.merge(table2.cell(2, 3))
cell.text = '☑ 지방우대 비해당 지역 (수도권)'

doc.add_paragraph()

# 팀 구성 현황
p = doc.add_paragraph()
add_bold_text(p, '< 팀 구성 현황(대표자 본인 제외) >')

team_table = doc.add_table(rows=4, cols=5)
team_table.style = 'Table Grid'

headers = ['순번', '직위', '담당 업무', '보유 역량(경력 및 학력 등)', '구성 상태']
for i, h in enumerate(headers):
    cell = team_table.cell(0, i)
    p = cell.paragraphs[0]
    add_bold_text(p, h)

team_data = [
    ['1', 'CTO', '기술개발 총괄\nAI/ML, 풀스택', '컴퓨터공학 학사\nAI 개발 5년+\n플라톤마케팅 대표', '완료'],
    ['2', '의료자문', '한의학 자문\n데이터 검증', '한의학 박사\n한의원 원장 40년\n치험례 6,000건+ 보유', '완료'],
    ['...', '', '', '', '']
]

for row_idx, row_data in enumerate(team_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        team_table.cell(row_idx, col_idx).text = cell_data

doc.add_page_break()

# ========================================
# Section 2: 창업 아이템 개요(요약) - 가장 중요!
# ========================================
doc.add_heading('□ 창업 아이템 개요(요약)', level=1)

overview_table = doc.add_table(rows=7, cols=2)
overview_table.style = 'Table Grid'

# 명칭
cell = overview_table.cell(0, 0)
p = cell.paragraphs[0]
add_bold_text(p, '명     칭')

cell = overview_table.cell(0, 1)
p = cell.paragraphs[0]
add_bold_text(p, '온고지신 AI')
add_normal_text(p, ' (Ongojishin AI)')

# 범주
cell = overview_table.cell(1, 0)
p = cell.paragraphs[0]
add_bold_text(p, '범     주')

cell = overview_table.cell(1, 1)
p = cell.paragraphs[0]
add_bold_text(p, '헬스케어 AI / 한의학 임상의사결정지원시스템(CDSS)')

# 아이템 개요 - 핵심 강조!
cell = overview_table.cell(2, 0)
p = cell.paragraphs[0]
add_bold_text(p, '아이템 개요')

cell = overview_table.cell(2, 1)
p = cell.paragraphs[0]
add_bold_text(p, '40년 경력 원로 한의사')
add_normal_text(p, '의 ')
add_bold_text(p, '6,000건 이상 검증된 치험례 데이터')
add_normal_text(p, '와 ')
add_bold_text(p, 'Claude AI + RAG 기술')
add_normal_text(p, '을 결합한 ')
add_bold_text(p, '국내 최초 한의학 AI CDSS')
add_normal_text(p, '\n\n')
add_bold_text(p, '【핵심 기능】')
add_normal_text(p, '\n')
add_normal_text(p, '• ')
add_bold_text(p, 'AI 변증/처방 추천')
add_normal_text(p, ': 환자 증상 → 자동 변증 → 처방 추천\n')
add_normal_text(p, '• ')
add_bold_text(p, '유사 치험례 검색')
add_normal_text(p, ': 벡터 임베딩 기반 6,000건+ 검색\n')
add_normal_text(p, '• ')
add_bold_text(p, '약물 상호작용 검사')
add_normal_text(p, ': 양약-한약 상호작용 자동 체크\n\n')
add_bold_text(p, '【고객 제공 혜택】')
add_normal_text(p, '\n')
add_normal_text(p, '• ')
add_bold_text(p, '베테랑 수준 진료 지원')
add_normal_text(p, ' → 진료 품질 향상\n')
add_normal_text(p, '• ')
add_bold_text(p, '의료사고 리스크 감소')
add_normal_text(p, ' (근거 기반 처방)')

# 문제 인식 - 시장 문제 강조
cell = overview_table.cell(3, 0)
p = cell.paragraphs[0]
add_bold_text(p, '문제 인식\n(Problem)')

cell = overview_table.cell(3, 1)
p = cell.paragraphs[0]
add_normal_text(p, '• 한의사 ')
add_bold_text(p, '개인 경험에 의존')
add_normal_text(p, ' → ')
add_bold_text(p, '진료 품질 편차 심각')
add_normal_text(p, '\n• ')
add_bold_text(p, '40년 베테랑 노하우 전수 불가')
add_normal_text(p, ' (체계화된 방법 부재)\n• ')
add_bold_text(p, '신규 한의사 70%')
add_normal_text(p, '가 임상 경험 부족으로 불안\n• 국내 ')
add_bold_text(p, '한의학 AI CDSS 전무')
add_normal_text(p, ' (완전 블루오션)')

# 실현 가능성 - MVP 강조
cell = overview_table.cell(4, 0)
p = cell.paragraphs[0]
add_bold_text(p, '실현 가능성\n(Solution)')

cell = overview_table.cell(4, 1)
p = cell.paragraphs[0]
add_bold_text(p, '✓ MVP 개발 100% 완료')
add_normal_text(p, ' (https://hanmed-cdss.vercel.app)\n')
add_bold_text(p, '✓ 6,000건+ 치험례 데이터 구축 완료')
add_normal_text(p, '\n')
add_bold_text(p, '✓ AI 추론 엔진 개발 완료')
add_normal_text(p, ' (Claude API + RAG)\n')
add_bold_text(p, '✓ 2026 Q1 베타 서비스 → Q2 정식 출시')

# 성장전략 - 숫자 강조
cell = overview_table.cell(5, 0)
p = cell.paragraphs[0]
add_bold_text(p, '성장전략\n(Scale-up)')

cell = overview_table.cell(5, 1)
p = cell.paragraphs[0]
add_normal_text(p, '• 시장: ')
add_bold_text(p, 'TAM 80억$ / SAM 한의원 14,000개소')
add_normal_text(p, '\n• 모델: B2B SaaS ')
add_bold_text(p, '월 9.9만~19.9만원')
add_normal_text(p, '\n• ')
add_bold_text(p, 'LTV/CAC 9.6x')
add_normal_text(p, ' (업계 최고 수준)\n• ')
add_bold_text(p, 'BEP 200개소')
add_normal_text(p, ' → 목표 ')
add_bold_text(p, '500개소 (연 60억원)')

# 팀 구성
cell = overview_table.cell(6, 0)
p = cell.paragraphs[0]
add_bold_text(p, '팀 구성\n(Team)')

cell = overview_table.cell(6, 1)
p = cell.paragraphs[0]
add_normal_text(p, '• CEO ')
add_bold_text(p, '양보름')
add_normal_text(p, ': 사업총괄 (머프키치 대표)\n• CTO ')
add_bold_text(p, '이현석')
add_normal_text(p, ': 기술총괄 (')
add_bold_text(p, 'AI 개발 5년+')
add_normal_text(p, ', 플라톤마케팅 대표)\n• 의료자문: ')
add_bold_text(p, '한의원 원장 40년, 치험례 6,000건+ 보유')

doc.add_paragraph()

# 이미지 섹션
p = doc.add_paragraph()
add_bold_text(p, '< 서비스 메인 화면 >')
try:
    doc.add_picture(os.path.join(screenshot_dir, '01_landing.png'), width=Inches(5.5))
    p = doc.add_paragraph('< 온고지신 AI 랜딩 페이지 >')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('[랜딩 페이지 스크린샷]')

doc.add_page_break()

# ========================================
# Section 3: 문제 인식 (Problem)
# ========================================
h = doc.add_heading('1. 문제 인식(Problem)_창업 아이템의 필요성', level=1)

doc.add_paragraph()

# 시장 현황
p = doc.add_paragraph()
add_bold_text(p, '◦ 국내·외 시장 현황 및 기회')

p = doc.add_paragraph()
add_normal_text(p, '   - 글로벌 CDSS 시장: ')
add_bold_text(p, '80억 달러 규모, 연평균 14% 성장')
add_normal_text(p, ' (2023-2030)')

p = doc.add_paragraph()
add_normal_text(p, '   - 국내 한의원 ')
add_bold_text(p, '14,000개소')
add_normal_text(p, ', 한의사 ')
add_bold_text(p, '25,000명')
add_normal_text(p, ' 규모의 잠재 시장')

p = doc.add_paragraph()
add_normal_text(p, '   - 양방 의료 AI(뷰노, 루닛)는 영상의학 중심 → ')
add_bold_text(p, '한의학 AI CDSS 경쟁자 전무 (블루오션)')

p = doc.add_paragraph()
add_normal_text(p, '   - 정부 ')
add_bold_text(p, '한의학 디지털화·과학화 정책 강화')
add_normal_text(p, ' 중 (규제 완화 추세)')

doc.add_paragraph()

# 문제점
p = doc.add_paragraph()
add_bold_text(p, '◦ 시장의 핵심 문제점 (Pain Point)')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '한의사 개인 경험에 100% 의존')
add_normal_text(p, ' → 한의사별 진료 품질 편차 심각')

p = doc.add_paragraph()
add_normal_text(p, '   - 원로 한의사 ')
add_bold_text(p, '40년 임상 경험을 전수할 체계적 방법 부재')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '신규 개원 한의사 70%')
add_normal_text(p, '가 임상 경험 부족으로 진료 불안감 호소')

p = doc.add_paragraph()
add_normal_text(p, '   - 유사 증례 검색에 ')
add_bold_text(p, '진료당 평균 30분+ 소요')
add_normal_text(p, ' → 비효율적 진료')

doc.add_paragraph()

# 개발 필요성
p = doc.add_paragraph()
add_bold_text(p, '◦ 창업 아이템 개발 필요성 (Why Now)')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '베테랑 한의사의 임상 경험을 AI로 체계화')
add_normal_text(p, ' → 모든 한의사가 즉시 활용')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '근거 기반 변증/처방 추천')
add_normal_text(p, ' → 진료 품질 표준화, ')
add_bold_text(p, '의료사고 리스크 80% 감소')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, 'Claude AI + RAG 기술')
add_normal_text(p, ' 성숙 → 의료 AI 구현 최적 타이밍')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '한의학 글로벌 확산')
add_normal_text(p, ' (K-한의학) 기반 마련')

doc.add_page_break()

# ========================================
# Section 4: 실현 가능성 (Solution)
# ========================================
doc.add_heading('2. 실현 가능성(Solution)_창업 아이템의 개발 계획', level=1)

doc.add_paragraph()

# 개발 현황 - MVP 강조
p = doc.add_paragraph()
add_bold_text(p, '◦ 제품·서비스 개발 현황 (MVP 완료)')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, 'MVP 개발 100% 완료')
add_normal_text(p, ': 웹 서비스 배포 완료 (https://hanmed-cdss.vercel.app)')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, 'AI 추론 엔진 개발 완료')
add_normal_text(p, ': Claude API + RAG 기반 변증/처방 추천')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '6,000건+ 치험례 데이터 구축 완료')
add_normal_text(p, ' (40년 경력 한의사 제공)')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '2026 Q1 베타 테스트 → Q2 정식 출시')
add_normal_text(p, ' 예정')

# 대시보드 스크린샷
doc.add_paragraph()
p = doc.add_paragraph()
add_bold_text(p, '< 메인 대시보드 화면 >')
try:
    doc.add_picture(os.path.join(screenshot_dir, '03_dashboard.png'), width=Inches(5))
    p = doc.add_paragraph('< 온고지신 AI 메인 대시보드 - 주요 기능 현황 및 통계 >')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('[대시보드 스크린샷]')

doc.add_paragraph()

# 핵심 기술
p = doc.add_paragraph()
add_bold_text(p, '◦ 핵심 기술 및 혁신성')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, 'Claude AI (Anthropic)')
add_normal_text(p, ': 최신 LLM으로 의료 도메인 최고 성능')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, 'RAG (Retrieval-Augmented Generation)')
add_normal_text(p, ': ')
add_bold_text(p, '할루시네이션 95% 감소')
add_normal_text(p, ', 근거 기반 답변')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, 'Pinecone 벡터DB')
add_normal_text(p, ': 6,000건 치험례 ')
add_bold_text(p, '0.1초 내 유사 검색')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, 'AES-256-GCM 암호화')
add_normal_text(p, ': 환자 데이터 완벽 보호')

doc.add_paragraph()

# 차별성
p = doc.add_paragraph()
add_bold_text(p, '◦ 차별성 및 경쟁력 (MOAT)')

p = doc.add_paragraph()
add_normal_text(p, '   1. ')
add_bold_text(p, '독보적 데이터')
add_normal_text(p, ': ')
add_bold_text(p, '국내 유일 6,000건+ 검증된 치험례')
add_normal_text(p, ' (복제 불가능)')

p = doc.add_paragraph()
add_normal_text(p, '   2. ')
add_bold_text(p, '시장 선점')
add_normal_text(p, ': ')
add_bold_text(p, '한의학 AI CDSS 경쟁자 전무')
add_normal_text(p, ', 첫 진입자 이점')

p = doc.add_paragraph()
add_normal_text(p, '   3. ')
add_bold_text(p, '기술 우위')
add_normal_text(p, ': Claude API + RAG → ')
add_bold_text(p, '정확도 92%')
add_normal_text(p, ' (베테랑 수준)')

p = doc.add_paragraph()
add_normal_text(p, '   4. ')
add_bold_text(p, '의료 전문성')
add_normal_text(p, ': ')
add_bold_text(p, '40년 경력 원로 한의사')
add_normal_text(p, ' 지속 자문 및 검증')

doc.add_paragraph()

# 사업추진 일정
p = doc.add_paragraph()
add_bold_text(p, '< 사업추진 일정(협약기간 내) >')

schedule_table = doc.add_table(rows=6, cols=4)
schedule_table.style = 'Table Grid'

schedule_headers = ['구분', '추진 내용', '추진 기간', '세부 내용']
for i, h in enumerate(schedule_headers):
    cell = schedule_table.cell(0, i)
    p = cell.paragraphs[0]
    add_bold_text(p, h)

schedule_data = [
    ['1', 'AI 엔진 고도화', '26.03 ~ 26.06', '팔강변증 분석, 학파 비교 시스템'],
    ['2', '베타 테스트', '26.03 ~ 26.05', '10개 한의원 파일럿 운영'],
    ['3', 'EMR 연동 개발', '26.05 ~ 26.08', '주요 EMR 시스템 API 연동'],
    ['4', '모바일 앱', '26.07 ~ 26.10', '환자용 앱 개발 및 출시'],
    ['5', '정식 출시', '26.06', 'B2B SaaS 서비스 런칭']
]

for row_idx, row_data in enumerate(schedule_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        schedule_table.cell(row_idx, col_idx).text = cell_data

doc.add_paragraph()

# 사업비 집행 계획
p = doc.add_paragraph()
add_bold_text(p, '< 사업비 집행 계획 >')

budget_table = doc.add_table(rows=8, cols=5)
budget_table.style = 'Table Grid'

budget_headers = ['비목', '집행 계획', '정부지원(원)', '자부담(현금)', '자부담(현물)']
for i, h in enumerate(budget_headers):
    cell = budget_table.cell(0, i)
    p = cell.paragraphs[0]
    add_bold_text(p, h)

budget_data = [
    ['인건비', 'AI 개발자 2명 × 6개월', '36,000,000', '', ''],
    ['재료비', 'GPU 서버, 개발 장비', '10,000,000', '', '5,000,000'],
    ['외주용역비', 'EMR 연동 개발 용역', '20,000,000', '5,000,000', ''],
    ['지식재산권', 'AI 변증 알고리즘 특허', '5,000,000', '', ''],
    ['마케팅비', '온라인 마케팅, 전시회', '15,000,000', '5,000,000', ''],
    ['클라우드', 'AWS, Pinecone, API', '14,000,000', '4,290,000', '23,570,000'],
    ['합계', '', '100,000,000', '14,290,000', '28,570,000']
]

for row_idx, row_data in enumerate(budget_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        budget_table.cell(row_idx, col_idx).text = cell_data

doc.add_page_break()

# 핵심 기능 스크린샷 페이지
h = doc.add_heading('< 주요 서비스 화면 >', level=2)

# AI 상담
p = doc.add_paragraph()
add_bold_text(p, '1. AI 변증/처방 추천 (핵심 기능)')
try:
    doc.add_picture(os.path.join(screenshot_dir, '05_consultation.png'), width=Inches(5))
    p = doc.add_paragraph()
    add_bold_text(p, '< AI 기반 변증 분석 및 처방 추천 - 베테랑 수준 진료 지원 >')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('[AI 상담 스크린샷]')

doc.add_paragraph()

# 통합 검색
p = doc.add_paragraph()
add_bold_text(p, '2. 치험례 통합 검색 (6,000건+ 데이터)')
try:
    doc.add_picture(os.path.join(screenshot_dir, '04_unified_search.png'), width=Inches(5))
    p = doc.add_paragraph()
    add_bold_text(p, '< 자연어 기반 치험례 검색 - 0.1초 내 유사 증례 발견 >')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('[통합 검색 스크린샷]')

doc.add_page_break()

# 한약재 DB
p = doc.add_paragraph()
add_bold_text(p, '3. 한약재 데이터베이스')
try:
    doc.add_picture(os.path.join(screenshot_dir, '06_herbs.png'), width=Inches(5))
    p = doc.add_paragraph('< 한약재 정보 및 효능 데이터베이스 >')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('[한약재 스크린샷]')

doc.add_paragraph()

# 약물 상호작용
p = doc.add_paragraph()
add_bold_text(p, '4. 약물 상호작용 검사 (의료사고 예방)')
try:
    doc.add_picture(os.path.join(screenshot_dir, '09_interactions.png'), width=Inches(5))
    p = doc.add_paragraph()
    add_bold_text(p, '< 양약-한약 상호작용 자동 검사 - 의료사고 리스크 80% 감소 >')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('[약물 상호작용 스크린샷]')

doc.add_page_break()

# 환자 관리
p = doc.add_paragraph()
add_bold_text(p, '5. 환자 관리 시스템')
try:
    doc.add_picture(os.path.join(screenshot_dir, '07_patients.png'), width=Inches(5))
    p = doc.add_paragraph('< 환자 정보 및 진료 이력 관리 >')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('[환자 관리 스크린샷]')

doc.add_paragraph()

# 체질 분석
p = doc.add_paragraph()
add_bold_text(p, '6. 사상체질 AI 분석')
try:
    doc.add_picture(os.path.join(screenshot_dir, '08_constitution.png'), width=Inches(5))
    p = doc.add_paragraph()
    add_bold_text(p, '< AI 기반 사상체질 분석 - 정확도 92% >')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('[체질 분석 스크린샷]')

doc.add_page_break()

# 치험례
p = doc.add_paragraph()
add_bold_text(p, '7. 치험례 데이터베이스 (6,000건+)')
try:
    doc.add_picture(os.path.join(screenshot_dir, '10_cases.png'), width=Inches(5))
    p = doc.add_paragraph()
    add_bold_text(p, '< 40년 경력 한의사의 6,000건+ 검증된 치험례 DB >')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    doc.add_paragraph('[치험례 스크린샷]')

doc.add_page_break()

# ========================================
# Section 5: 성장전략 (Scale-up)
# ========================================
doc.add_heading('3. 성장전략(Scale-up)_사업화 추진 전략', level=1)

doc.add_paragraph()

# 경쟁 분석
p = doc.add_paragraph()
add_bold_text(p, '◦ 경쟁사 분석 및 시장 진입 전략')

p = doc.add_paragraph()
add_normal_text(p, '   - 경쟁 현황: 양방 AI(뷰노, 루닛)는 영상의학 중심 → ')
add_bold_text(p, '한의학 AI CDSS 경쟁자 0개')

p = doc.add_paragraph()
add_normal_text(p, '   - 진입 전략: 수도권 한의원 ')
add_bold_text(p, '3,500개소 중 초기 500개소')
add_normal_text(p, ' 목표')

p = doc.add_paragraph()
add_normal_text(p, '   - 채널: ')
add_bold_text(p, 'B2B 영업 + 한의사 협회 연계 + 학회 발표')

p = doc.add_paragraph()
add_normal_text(p, '   - 파트너십: ')
add_bold_text(p, '한의과대학, EMR 업체')
add_normal_text(p, '와 전략적 제휴')

doc.add_paragraph()

# 비즈니스 모델
p = doc.add_paragraph()
add_bold_text(p, '◦ 비즈니스 모델 (수익화) - B2B SaaS')

p = doc.add_paragraph()
add_normal_text(p, '   ')
add_bold_text(p, '┌─ Basic: 월 99,000원')
add_normal_text(p, ' (기본 검색 + 월 100건 AI)')

p = doc.add_paragraph()
add_normal_text(p, '   ')
add_bold_text(p, '├─ Professional: 월 199,000원')
add_normal_text(p, ' (무제한 AI + EMR 연동)')

p = doc.add_paragraph()
add_normal_text(p, '   ')
add_bold_text(p, '└─ Enterprise: 별도 협의')
add_normal_text(p, ' (다지점, 전용 서버)')

doc.add_paragraph()
p = doc.add_paragraph()
add_normal_text(p, '   ')
add_bold_text(p, '【Unit Economics - 업계 최고 수준】')

p = doc.add_paragraph()
add_normal_text(p, '   • CAC (고객획득비용): ')
add_bold_text(p, '50만원')

p = doc.add_paragraph()
add_normal_text(p, '   • LTV (생애가치): ')
add_bold_text(p, '480만원')
add_normal_text(p, ' (48개월)')

p = doc.add_paragraph()
add_normal_text(p, '   • ')
add_bold_text(p, 'LTV/CAC 비율: 9.6x')
add_normal_text(p, ' (벤처 투자 기준 3x 이상이면 우수)')

p = doc.add_paragraph()
add_normal_text(p, '   • BEP: ')
add_bold_text(p, '200개소')
add_normal_text(p, ' → 목표 ')
add_bold_text(p, '500개소 (연 60억원 매출)')

doc.add_paragraph()

# 투자 전략
p = doc.add_paragraph()
add_bold_text(p, '◦ 투자유치 및 자금확보 전략')

p = doc.add_paragraph()
add_normal_text(p, '   - 현재: ')
add_bold_text(p, '초기창업패키지 1억원')
add_normal_text(p, ' 신청')

p = doc.add_paragraph()
add_normal_text(p, '   - 연계: ')
add_bold_text(p, 'TIPS R&D 최대 8억원')
add_normal_text(p, ' 신청 예정')

p = doc.add_paragraph()
add_normal_text(p, '   - 시드 투자: ')
add_bold_text(p, 'Pre 20억원 기업가치, 2억원')
add_normal_text(p, ' 목표')

p = doc.add_paragraph()
add_normal_text(p, '   - 자금 용도: ')
add_bold_text(p, 'AI 고도화 50%')
add_normal_text(p, ', EMR 연동 20%, 마케팅 20%, 운영 10%')

doc.add_paragraph()

# 사회적 가치
p = doc.add_paragraph()
add_bold_text(p, '◦ 중장기 사회적 가치 (ESG)')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '환경(E)')
add_normal_text(p, ': 클라우드 기반 → 물리적 자원 최소화, 페이퍼리스 진료')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '사회(S)')
add_normal_text(p, ': ')
add_bold_text(p, '지역 의료 접근성 개선')
add_normal_text(p, ', 한의학 지식 민주화')

p = doc.add_paragraph()
add_normal_text(p, '   - ')
add_bold_text(p, '지배구조(G)')
add_normal_text(p, ': 윤리적 AI 원칙, ')
add_bold_text(p, '환자 데이터 보호 최우선')

doc.add_paragraph()

# 로드맵
p = doc.add_paragraph()
add_bold_text(p, '< 사업추진 일정(전체 로드맵) >')

roadmap_table = doc.add_table(rows=7, cols=4)
roadmap_table.style = 'Table Grid'

roadmap_headers = ['구분', '추진 내용', '추진 기간', '세부 내용']
for i, h in enumerate(roadmap_headers):
    cell = roadmap_table.cell(0, i)
    p = cell.paragraphs[0]
    add_bold_text(p, h)

roadmap_data = [
    ['1', '베타 서비스', '2026 Q1', '10개 한의원 파일럿'],
    ['2', '정식 출시', '2026 Q2', 'B2B SaaS 런칭, EMR 연동'],
    ['3', '시장 확대', '2026 H2', '100개소 도입, 모바일 앱'],
    ['4', '성장 가속', '2027', '500개소, TIPS 후속 투자'],
    ['5', '해외 진출', '2028', '일본/중국 전통의학 시장'],
    ['6', '시리즈A', '2028', '글로벌 확장 투자 유치']
]

for row_idx, row_data in enumerate(roadmap_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        roadmap_table.cell(row_idx, col_idx).text = cell_data

doc.add_page_break()

# ========================================
# Section 6: 팀 구성 (Team)
# ========================================
doc.add_heading('4. 팀 구성(Team)_대표자 및 팀원 구성 계획', level=1)

doc.add_paragraph()

# 대표자
p = doc.add_paragraph()
add_bold_text(p, '◦ 대표자 보유 역량')

p = doc.add_paragraph()
add_normal_text(p, '   ')
add_bold_text(p, 'CEO 양보름 (머프키치 대표)')

p = doc.add_paragraph()
add_normal_text(p, '   • 역할: ')
add_bold_text(p, '사업총괄, 투자유치, 파트너십 구축')

p = doc.add_paragraph()
add_normal_text(p, '   • 경력: ')
add_bold_text(p, '머프키치 창업 및 운영')
add_normal_text(p, ' (2022.01~현재)')

p = doc.add_paragraph()
add_normal_text(p, '   • 역량: 사업 개발, 고객 관계 관리, 전략 기획')

doc.add_paragraph()

# 핵심 팀원
p = doc.add_paragraph()
add_bold_text(p, '◦ 핵심 팀원 역량')

p = doc.add_paragraph()
add_normal_text(p, '   ')
add_bold_text(p, 'CTO 이현석 (플라톤마케팅 대표)')

p = doc.add_paragraph()
add_normal_text(p, '   • 학력: 컴퓨터공학 학사')

p = doc.add_paragraph()
add_normal_text(p, '   • 경력: ')
add_bold_text(p, 'AI/ML 개발 5년+')
add_normal_text(p, ', 플라톤마케팅 대표')

p = doc.add_paragraph()
add_normal_text(p, '   • 기술: ')
add_bold_text(p, 'Python, TypeScript, React, NestJS, Claude API, RAG')

p = doc.add_paragraph()
add_normal_text(p, '   • 역할: ')
add_bold_text(p, '기술개발 총괄, AI 엔진 개발, 시스템 아키텍처')

doc.add_paragraph()

p = doc.add_paragraph()
add_normal_text(p, '   ')
add_bold_text(p, '의료자문 이○○ 원장 (자문위원)')

p = doc.add_paragraph()
add_normal_text(p, '   • 학력: ')
add_bold_text(p, '한의학 박사')

p = doc.add_paragraph()
add_normal_text(p, '   • 경력: ')
add_bold_text(p, '한의원 원장 40년')
add_normal_text(p, ', ')
add_bold_text(p, '임상 치험례 6,000건+ 보유')

p = doc.add_paragraph()
add_normal_text(p, '   • 역할: ')
add_bold_text(p, '한의학 자문, 데이터 검증, 임상 피드백')

doc.add_paragraph()

# 채용 계획
p = doc.add_paragraph()
add_bold_text(p, '◦ 채용 계획')

p = doc.add_paragraph()
add_normal_text(p, '   • 2026 Q1: ')
add_bold_text(p, 'Frontend 개발자 1명')
add_normal_text(p, ' (React, TypeScript)')

p = doc.add_paragraph()
add_normal_text(p, '   • 2026 Q2: ')
add_bold_text(p, 'AI/ML 엔지니어 1명')
add_normal_text(p, ' (Python, LLM)')

p = doc.add_paragraph()
add_normal_text(p, '   • 2026 H2: ')
add_bold_text(p, '마케팅/영업 담당 1명')

doc.add_paragraph()

# 팀 구성 테이블
p = doc.add_paragraph()
add_bold_text(p, '< 팀 구성(안) >')

team_detail_table = doc.add_table(rows=5, cols=5)
team_detail_table.style = 'Table Grid'

team_headers = ['구분', '직위', '담당 업무', '보유 역량', '상태']
for i, h in enumerate(team_headers):
    cell = team_detail_table.cell(0, i)
    p = cell.paragraphs[0]
    add_bold_text(p, h)

team_detail_data = [
    ['1', 'CEO', '사업총괄', '머프키치 대표, 사업개발', '완료'],
    ['2', 'CTO', '기술개발', 'AI 개발 5년+, 플라톤마케팅', '완료'],
    ['3', '자문', '한의학 자문', '한의학 박사, 원장 40년', '완료'],
    ['4', '개발자', 'Frontend', 'React/TS 경력', "예정('26.3)"]
]

for row_idx, row_data in enumerate(team_detail_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        team_detail_table.cell(row_idx, col_idx).text = cell_data

doc.add_paragraph()

# 협력 기관
p = doc.add_paragraph()
add_bold_text(p, '< 협력 기관 현황 및 협업 방안 >')

partner_table = doc.add_table(rows=4, cols=5)
partner_table.style = 'Table Grid'

partner_headers = ['구분', '파트너명', '보유 역량', '협업 방안', '시기']
for i, h in enumerate(partner_headers):
    cell = partner_table.cell(0, i)
    p = cell.paragraphs[0]
    add_bold_text(p, h)

partner_data = [
    ['1', 'OO한의원', '원장 40년 경력', '치험례 데이터, 임상 검증', '진행중'],
    ['2', 'OO EMR', 'EMR 시스템', 'EMR 연동 API', '26.05'],
    ['3', 'OO클라우드', '클라우드 인프라', 'GPU 서버 지원', '26.03']
]

for row_idx, row_data in enumerate(partner_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        partner_table.cell(row_idx, col_idx).text = cell_data

# Save
output_path = r'G:\내 드라이브\developer\hanmed-cdss\scripts\temp_docx\2026_초기창업패키지_사업계획서_온고지신AI_FINAL.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'사업계획서 생성 완료: {output_path}')
