# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)

# Title
title = doc.add_heading('초기창업패키지 창업기업 사업계획서', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Notice
notice = doc.add_paragraph()
run = notice.add_run('※ 사업계획서는 목차(1페이지)를 제외하고 15페이지 이내로 작성(증빙서류는 제한 없음)\n※ 본문 내 파란색 안내 문구는 삭제하고 검정 글씨로 작성하여 제출')
run.font.size = Pt(9)
notice.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

# Section 1: 일반현황
doc.add_heading('□ 일반현황', level=1)

# Table for 일반현황
table1 = doc.add_table(rows=6, cols=4)
table1.style = 'Table Grid'

# Row 1: 기업명
table1.cell(0, 0).text = '기업명'
table1.cell(0, 1).text = '(주)온고지신AI'
table1.cell(0, 2).text = '개업연월일'
table1.cell(0, 3).text = 'OO.OO.OO'

# Row 2: 사업자 구분
table1.cell(1, 0).text = '사업자 구분\n(모집마감일 기준)'
table1.cell(1, 1).text = '법인사업자'
table1.cell(1, 2).text = '대표자 유형\n(모집마감일 기준)'
table1.cell(1, 3).text = '공동대표'

# Row 3: 사업자등록번호
table1.cell(2, 0).text = '사업자등록번호\n(법인등록번호)'
table1.cell(2, 1).text = 'OOO-OO-OOOOO\n(OOOOOO-OOOOOOO)'
table1.cell(2, 2).text = '사업장 소재지\n(본점)'
table1.cell(2, 3).text = 'OO도 OO시·군'

# Row 4: 창업아이템명
table1.cell(3, 0).text = '창업아이템명'
cell = table1.cell(3, 1)
cell.merge(table1.cell(3, 3))
cell.text = 'Claude AI + RAG 기술이 적용된 한의학 변증/처방 추천 기능의 한의사 진료 품질 향상을 제공하는 임상의사결정지원시스템(CDSS) 웹 서비스'

# Row 5: 산출물
table1.cell(4, 0).text = '산출물\n(협약기간 내 목표)'
cell = table1.cell(4, 1)
cell.merge(table1.cell(4, 3))
cell.text = '웹사이트(1개), AI 추론 엔진(1개), 모바일 앱(1개 - 환자용)'

# Row 6: 지원분야
table1.cell(5, 0).text = '지원 분야 (택 1)'
cell = table1.cell(5, 1)
cell.merge(table1.cell(5, 3))
cell.text = '☑ 지식서비스   ☐ 제조'

doc.add_paragraph()

# 추가 일반현황 테이블
table2 = doc.add_table(rows=3, cols=4)
table2.style = 'Table Grid'

table2.cell(0, 0).text = '전문기술분야 (택 1)'
cell = table2.cell(0, 1)
cell.merge(table2.cell(0, 3))
cell.text = '☑ 정보·통신 (인공지능)   ☐ 바이오·의료·생명   ☐ 기타'

table2.cell(1, 0).text = '총 사업비\n구성 계획'
table2.cell(1, 1).text = '정부지원사업비(A)\n100백만원'
table2.cell(1, 2).text = '자기부담사업비(B)\n현금: 14.29백만원\n현물: 28.57백만원'
table2.cell(1, 3).text = '총 사업비(C=A+B)\n142.86백만원'

table2.cell(2, 0).text = '지방우대 지역\n해당여부'
cell = table2.cell(2, 1)
cell.merge(table2.cell(2, 3))
cell.text = '☐ 특별지원 지역  ☐ 우대지원 지역  ☐ 일반지역  ☑ 지방우대 비해당 지역'

doc.add_paragraph()

# 팀 구성 현황 table
doc.add_paragraph('< 팀 구성 현황(대표자 본인 제외) >')
team_table = doc.add_table(rows=5, cols=5)
team_table.style = 'Table Grid'

headers = ['순번', '직위', '담당 업무', '보유 역량(경력 및 학력 등)', '구성 상태']
for i, h in enumerate(headers):
    team_table.cell(0, i).text = h

team_data = [
    ['1', 'CTO', '기술개발 총괄\nAI/ML, 풀스택 개발', '컴퓨터공학 학사, AI/ML 개발 경력 5년+', '완료'],
    ['2', '의료자문', '한의학 자문\n데이터 검증', '한의학 박사, 한의원 원장 40년 경력', '완료'],
    ['3', '개발자', 'Frontend 개발', 'OO학 학사, React 개발 경력', "예정('26.3)"],
    ['...', '', '', '', '']
]

for row_idx, row_data in enumerate(team_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        team_table.cell(row_idx, col_idx).text = cell_data

doc.add_page_break()

# Section 2: 창업 아이템 개요(요약)
doc.add_heading('□ 창업 아이템 개요(요약)', level=1)

overview_table = doc.add_table(rows=7, cols=2)
overview_table.style = 'Table Grid'

# 명칭
overview_table.cell(0, 0).text = '명     칭'
overview_table.cell(0, 1).text = '온고지신 AI (Ongojishin AI)'

# 범주
overview_table.cell(1, 0).text = '범     주'
overview_table.cell(1, 1).text = '헬스케어 AI / 의료 임상의사결정지원시스템(CDSS)'

# 아이템 개요
overview_table.cell(2, 0).text = '아이템 개요'
overview_table.cell(2, 1).text = '''40년 경력 원로 한의사의 6,000건 이상 검증된 치험례 데이터와 Claude AI + RAG(Retrieval-Augmented Generation) 기술을 결합하여, 한의사들의 변증(辨證) 분석 및 처방 추천을 지원하는 B2B SaaS 플랫폼

【핵심 기능】
• AI 기반 변증/처방 추천: 환자 증상 입력 → 자동 변증 분석 → 처방 추천
• 유사 치험례 검색: 벡터 임베딩 기반 의미 검색으로 6,000건+ 사례 검색
• 약물 상호작용 검사: 양약-한약 상호작용 자동 체크
• 팔강변증 분석: 음양/표리/한열/허실 분석 및 신뢰도 점수 제시

【고객 제공 혜택】
• 베테랑 수준의 임상 의사결정 지원 → 진료 품질 향상
• 근거 기반 처방으로 의료 사고 리스크 감소
• 진료 시간 단축 및 효율성 증대'''

# 문제 인식
overview_table.cell(3, 0).text = '문제 인식\n(Problem)'
overview_table.cell(3, 1).text = '''• 한의사 개인 경험에 의존하는 진료 → 품질 편차 발생
• 40년 베테랑의 임상 경험을 체계화/전수할 방법 부재
• 신규 한의사의 경험 부족으로 인한 실전 불안감
• 유사 증례 검색에 많은 시간 소요'''

# 실현 가능성
overview_table.cell(4, 0).text = '실현 가능성\n(Solution)'
overview_table.cell(4, 1).text = '''• MVP 개발 완료 (https://hanmed-cdss.vercel.app)
• 6,000건 이상 검증된 치험례 데이터 구축 완료
• Claude AI + RAG 기반 AI 추론 엔진 개발 완료
• 2026 Q1 베타 서비스 오픈, Q2 정식 출시 예정'''

# 성장전략
overview_table.cell(5, 0).text = '성장전략\n(Scale-up)'
overview_table.cell(5, 1).text = '''• 시장 규모: TAM 80억$ + SAM 국내 한의원 14,000개소
• 비즈니스 모델: B2B SaaS 구독 (월 9.9만~19.9만원)
• LTV/CAC: 9.6x (우수한 유닛 이코노믹스)
• 투자 유치: TIPS R&D 연계 (최대 8억원) + 시드 투자 2억원'''

# 팀 구성
overview_table.cell(6, 0).text = '팀 구성\n(Team)'
overview_table.cell(6, 1).text = '''• CEO 양○○: 사업총괄, 전략, 영업
• CTO 이○○: 기술개발, AI/ML, 풀스택 개발
• 의료자문 이○○ 원장: 한의학 자문, 40년 임상 경험, 데이터 검증'''

# 이미지 섹션
doc.add_paragraph()
img_table = doc.add_table(rows=2, cols=2)
img_table.style = 'Table Grid'
img_table.cell(0, 0).text = '[서비스 화면 스크린샷 삽입]'
img_table.cell(0, 1).text = '[AI 분석 결과 화면 삽입]'
img_table.cell(1, 0).text = '< 온고지신 AI 메인 대시보드 >'
img_table.cell(1, 1).text = '< AI 변증/처방 추천 화면 >'

doc.add_page_break()

# Section 3: 문제 인식 (Problem)
doc.add_heading('1. 문제 인식(Problem)_창업 아이템의 필요성', level=1)

doc.add_paragraph()
p1 = doc.add_paragraph()
p1.add_run('◦ 국내·외 시장 현황').bold = True

doc.add_paragraph('''   - 글로벌 CDSS 시장: 80억 달러 규모, 연평균 14% 성장 (2023-2030)
   - 국내 한의원 14,000개소, 한의사 25,000명 규모의 잠재 시장
   - 양방 의료 AI(뷰노, 루닛 등)는 영상의학 중심으로 성장 중
   - 한의학 분야 AI CDSS는 경쟁자가 전무한 블루오션''')

p2 = doc.add_paragraph()
p2.add_run('◦ 시장의 문제점').bold = True

doc.add_paragraph('''   - 한의학 진료는 개인 경험에 크게 의존 → 한의사별 진료 품질 편차 발생
   - 원로 한의사의 40년 임상 경험을 체계적으로 전수할 방법이 부재
   - 신규 개원 한의사의 임상 경험 부족으로 인한 불안감과 진료 품질 저하
   - 과거 유사 증례를 검색하는 데 많은 시간 소요, 비효율적 진료 프로세스''')

p3 = doc.add_paragraph()
p3.add_run('◦ 창업 아이템 개발 필요성').bold = True

doc.add_paragraph('''   - 베테랑 한의사의 임상 경험을 AI로 체계화하여 모든 한의사가 활용 가능하게 함
   - 근거 기반 변증/처방 추천으로 진료 품질 표준화 및 의료 사고 리스크 감소
   - 실시간 AI 지원으로 진료 효율성 증대 및 환자 만족도 향상
   - 한의학의 과학적 재조명 및 글로벌 확산 기반 마련''')

doc.add_page_break()

# Section 4: 실현 가능성 (Solution)
doc.add_heading('2. 실현 가능성(Solution)_창업 아이템의 개발 계획', level=1)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.add_run('◦ 제품·서비스 개발 현황 및 계획').bold = True

doc.add_paragraph('''   - MVP 개발 완료: 웹 서비스 (https://hanmed-cdss.vercel.app) 배포 완료
   - AI 추론 엔진: Claude API + RAG 기반 변증/처방 추천 시스템 구축 완료
   - 데이터 구축: 40년 경력 한의사의 6,000건+ 치험례 데이터 구조화 완료
   - 베타 테스트: 2026 Q1 10개 한의원 파일럿 운영 예정''')

p5 = doc.add_paragraph()
p5.add_run('◦ 핵심 기술 및 아키텍처').bold = True

doc.add_paragraph('''   - AI/ML: Claude API (Anthropic), OpenAI Embedding, Pinecone (벡터DB)
   - RAG(Retrieval-Augmented Generation): 할루시네이션 최소화, 근거 기반 답변 보장
   - Backend: NestJS + FastAPI, PostgreSQL, Redis (캐싱)
   - Frontend: React + TypeScript, Tailwind CSS
   - 보안: 환자 데이터 AES-256-GCM 암호화, JWT 인증, RBAC''')

p6 = doc.add_paragraph()
p6.add_run('◦ 차별성 및 경쟁력 확보 전략').bold = True

doc.add_paragraph('''   - 독보적 데이터: 국내 유일 6,000건+ 검증된 치험례 데이터 보유 (복제 불가능)
   - 시장 선점: 한의학 AI CDSS 경쟁자 전무, 첫 진입자 이점 확보
   - 최신 AI 기술: Claude API + RAG로 정확하고 신뢰할 수 있는 추천 제공
   - 의료 전문성: 40년 경력 원로 한의사의 지속적인 데이터 검증 및 자문''')

doc.add_paragraph()

# 사업추진 일정 테이블
doc.add_paragraph('< 사업추진 일정(협약기간 내) >')

schedule_table = doc.add_table(rows=6, cols=4)
schedule_table.style = 'Table Grid'

schedule_headers = ['구분', '추진 내용', '추진 기간', '세부 내용']
for i, h in enumerate(schedule_headers):
    schedule_table.cell(0, i).text = h

schedule_data = [
    ['1', 'AI 엔진 고도화', '26.03 ~ 26.06', '팔강변증 분석기, 학파 비교 시스템 완성'],
    ['2', '베타 테스트 진행', '26.03 ~ 26.05', '10개 한의원 파일럿 운영 및 피드백 수집'],
    ['3', 'EMR 연동 개발', '26.05 ~ 26.08', '주요 한의원 EMR 시스템 연동 API 개발'],
    ['4', '모바일 앱 개발', '26.07 ~ 26.10', '환자용 모바일 앱 개발 및 출시'],
    ['5', '정식 서비스 출시', '26.06', '정식 B2B SaaS 서비스 런칭']
]

for row_idx, row_data in enumerate(schedule_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        schedule_table.cell(row_idx, col_idx).text = cell_data

doc.add_paragraph()

# 사업비 집행 계획
doc.add_paragraph('< 사업비 집행 계획 >')

budget_table = doc.add_table(rows=8, cols=5)
budget_table.style = 'Table Grid'

budget_headers = ['비목', '집행 계획', '정부지원사업비(원)', '자기부담(현금)', '자기부담(현물)']
for i, h in enumerate(budget_headers):
    budget_table.cell(0, i).text = h

budget_data = [
    ['인건비', 'AI 개발자 2명 × 6개월', '36,000,000', '', ''],
    ['재료비', 'GPU 서버, 개발 장비 등', '10,000,000', '', '5,000,000'],
    ['외주용역비', 'EMR 연동 개발 용역', '20,000,000', '5,000,000', ''],
    ['지식재산권', '특허 출원 (AI 변증 알고리즘)', '5,000,000', '', ''],
    ['마케팅비', '온라인 마케팅, 전시회 참가', '15,000,000', '5,000,000', ''],
    ['클라우드비용', 'AWS, Pinecone, API 비용', '14,000,000', '4,290,000', '23,570,000'],
    ['합계', '', '100,000,000', '14,290,000', '28,570,000']
]

for row_idx, row_data in enumerate(budget_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        budget_table.cell(row_idx, col_idx).text = cell_data

doc.add_page_break()

# Section 5: 성장전략 (Scale-up)
doc.add_heading('3. 성장전략(Scale-up)_사업화 추진 전략', level=1)

doc.add_paragraph()
p7 = doc.add_paragraph()
p7.add_run('◦ 경쟁사 분석 및 시장 진입 전략').bold = True

doc.add_paragraph('''   - 경쟁 현황: 양방 의료 AI(뷰노, 루닛)는 영상의학 중심, 한의학 AI CDSS 경쟁자 전무
   - 진입 전략: 수도권 한의원 3,500개소 중 초기 500개소 목표 → B2B 영업 + 학회 연계
   - 파트너십: 한의사 협회, 한의과대학, EMR 업체와 전략적 제휴''')

p8 = doc.add_paragraph()
p8.add_run('◦ 비즈니스 모델 (수익화 모델)').bold = True

doc.add_paragraph('''   - B2B SaaS 구독 모델
     · Basic: 월 99,000원 (기본 검색 + 월 100건 AI 분석)
     · Professional: 월 199,000원 (무제한 AI + EMR 연동)
     · Enterprise: 별도 협의 (다지점, 전용 서버)
   - Unit Economics
     · CAC (고객획득비용): 50만원
     · LTV (생애가치): 480만원 (48개월)
     · LTV/CAC 비율: 9.6x (매우 우수)
     · BEP (손익분기점): 200개소''')

p9 = doc.add_paragraph()
p9.add_run('◦ 투자유치 및 자금확보 전략').bold = True

doc.add_paragraph('''   - 현재: 초기창업패키지 신청 (최대 1억원)
   - 연계: TIPS R&D 지원사업 신청 예정 (최대 8억원)
   - 시드 투자: Pre 20억원 기업가치, 2억원 투자 유치 목표
   - 자금 용도: AI 고도화 50%, EMR 연동 20%, 마케팅 20%, 운영 10%''')

p10 = doc.add_paragraph()
p10.add_run('◦ 중장기 사회적 가치 도입 계획').bold = True

doc.add_paragraph('''   - 환경: 클라우드 기반 서비스로 물리적 자원 사용 최소화, 페이퍼리스 진료 지원
   - 사회: 지역 의료 접근성 개선 (원격 진료 지원), 한의학 지식 민주화
   - 지배구조: 윤리적 AI 개발 원칙 준수, 환자 데이터 보호 최우선''')

doc.add_paragraph()

# 사업추진 일정(전체)
doc.add_paragraph('< 사업추진 일정(전체 사업단계) >')

roadmap_table = doc.add_table(rows=7, cols=4)
roadmap_table.style = 'Table Grid'

roadmap_headers = ['구분', '추진 내용', '추진 기간', '세부 내용']
for i, h in enumerate(roadmap_headers):
    roadmap_table.cell(0, i).text = h

roadmap_data = [
    ['1', '베타 서비스 오픈', '2026 Q1', '10개 한의원 파일럿, 피드백 수집'],
    ['2', '정식 서비스 출시', '2026 Q2', 'B2B SaaS 런칭, EMR 연동'],
    ['3', '시장 확대', '2026 H2', '100개소 도입, 모바일 앱 출시'],
    ['4', '성장 가속화', '2027', '500개소 달성, TIPS 후속 투자'],
    ['5', '해외 진출', '2028', '일본/중국 전통의학 시장 진출 검토'],
    ['6', '시리즈A 투자', '2028', '글로벌 확장을 위한 대규모 투자 유치']
]

for row_idx, row_data in enumerate(roadmap_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        roadmap_table.cell(row_idx, col_idx).text = cell_data

doc.add_page_break()

# Section 6: 팀 구성 (Team)
doc.add_heading('4. 팀 구성(Team)_대표자 및 팀원 구성 계획', level=1)

doc.add_paragraph()
p11 = doc.add_paragraph()
p11.add_run('◦ 대표자 보유 역량').bold = True

doc.add_paragraph('''   - CEO 양○○
     · 학력: OO대학교 경영학과 학사
     · 경력: 스타트업 창업 경험, 사업 개발 및 영업 전략 수립
     · 역할: 사업총괄, 투자유치, 파트너십 구축, 영업 전략''')

p12 = doc.add_paragraph()
p12.add_run('◦ 핵심 팀원 역량').bold = True

doc.add_paragraph('''   - CTO 이○○
     · 학력: OO대학교 컴퓨터공학과 학사
     · 경력: AI/ML 개발 5년+, 풀스택 개발 경험
     · 기술: Python, TypeScript, React, NestJS, Claude API, RAG
     · 역할: 기술개발 총괄, AI 엔진 개발, 시스템 아키텍처 설계

   - 의료자문 이○○ 원장 (자문위원)
     · 학력: OO대학교 한의학과 박사
     · 경력: 한의원 원장 40년, 임상 치험례 6,000건+ 보유
     · 역할: 한의학 자문, 데이터 검증, 임상 피드백 제공''')

p13 = doc.add_paragraph()
p13.add_run('◦ 채용 계획').bold = True

doc.add_paragraph('''   - 2026 Q1: Frontend 개발자 1명 (React, TypeScript)
   - 2026 Q2: AI/ML 엔지니어 1명 (Python, LLM)
   - 2026 H2: 마케팅/영업 담당 1명''')

doc.add_paragraph()

# 팀 구성 테이블
doc.add_paragraph('< 팀 구성(안) >')

team_detail_table = doc.add_table(rows=5, cols=5)
team_detail_table.style = 'Table Grid'

team_headers = ['구분', '직위', '담당 업무', '보유 역량(경력 및 학력 등)', '구성 상태']
for i, h in enumerate(team_headers):
    team_detail_table.cell(0, i).text = h

team_detail_data = [
    ['1', 'CEO', '사업총괄, 영업', '경영학 학사, 스타트업 창업 경험', '완료'],
    ['2', 'CTO', '기술개발 총괄', '컴퓨터공학 학사, AI/ML 개발 5년+', '완료'],
    ['3', '자문위원', '한의학 자문', '한의학 박사, 한의원 원장 40년', '완료'],
    ['4', '개발자', 'Frontend 개발', 'OO학 학사, React 경력', "예정('26.3)"]
]

for row_idx, row_data in enumerate(team_detail_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        team_detail_table.cell(row_idx, col_idx).text = cell_data

doc.add_paragraph()

# 협력 기관 현황
doc.add_paragraph('< 협력 기관 현황 및 협업 방안 >')

partner_table = doc.add_table(rows=4, cols=5)
partner_table.style = 'Table Grid'

partner_headers = ['구분', '파트너명', '보유 역량', '협업 방안', '협력 시기']
for i, h in enumerate(partner_headers):
    partner_table.cell(0, i).text = h

partner_data = [
    ['1', 'OO한의원', '한의원 원장 40년 경력', '치험례 데이터 제공, 임상 검증', '진행 중'],
    ['2', 'OO EMR 업체', 'EMR 시스템 개발/운영', 'EMR 연동 API 개발 협력', '26.05'],
    ['3', 'OO 클라우드', '클라우드 인프라', 'GPU 서버 및 스토리지 지원', '26.03']
]

for row_idx, row_data in enumerate(partner_data, 1):
    for col_idx, cell_data in enumerate(row_data):
        partner_table.cell(row_idx, col_idx).text = cell_data

# Save
output_path = r'G:\내 드라이브\developer\hanmed-cdss\scripts\temp_docx\2026_초기창업패키지_사업계획서_온고지신AI.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'사업계획서 생성 완료: {output_path}')
