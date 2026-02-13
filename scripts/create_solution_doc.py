"""
실현 가능성(Solution) - 창업 아이템의 개발 계획 워드 문서 생성
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 경로 설정
BASE_DIR = r"G:\내 드라이브\developer\hanmed-cdss"
SCREENSHOTS_DIR = os.path.join(BASE_DIR, "docs", "screenshots")
OUTPUT_PATH = os.path.join(BASE_DIR, "docs", "실현가능성_창업아이템_개발계획_v2.docx")

def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, bold=False, font_size=11, space_after=6):
    """포맷된 문단 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading_custom(doc, text, level=1):
    """커스텀 제목 추가"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return heading

def add_table_with_header(doc, headers, rows, col_widths=None):
    """헤더가 있는 테이블 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], 'E8F5E9')

    # 데이터 행
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    run.font.size = Pt(10)

    # 열 너비 설정
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def add_image_with_caption(doc, image_path, caption, width=5.5):
    """캡션이 있는 이미지 추가"""
    if os.path.exists(image_path):
        # 이미지 추가
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))

        # 캡션 추가
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run(caption)
        caption_run.font.size = Pt(9)
        caption_run.font.name = '맑은 고딕'
        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        caption_run.italic = True
        caption_p.paragraph_format.space_after = Pt(12)
    else:
        p = doc.add_paragraph(f"[이미지 없음: {caption}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_document():
    doc = Document()

    # 문서 여백 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== 제목 ==========
    title = doc.add_heading('실현 가능성(Solution)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('창업 아이템의 개발 계획')
    run.font.size = Pt(16)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_paragraph()

    # ========== 1. 아이템 개요 ==========
    add_heading_custom(doc, '1. 아이템 개요', level=1)

    # 1-1. 필요성 및 배경
    add_heading_custom(doc, '1-1. 필요성 및 배경', level=2)

    add_formatted_paragraph(doc,
        "우리나라는 고령화 사회로 빠르게 진입하면서 만성질환자가 급격히 증가하고 있음. "
        "보건복지부 통계에 따르면, 65세 이상 인구의 89.5%가 1개 이상의 만성질환을 보유하고 있으며, "
        "이에 따라 양방과 한방을 병행하는 환자도 늘어나고 있음. 정부는 '제4차 한의약육성발전종합계획(2026~2030)'을 "
        "통해 한의학의 과학화 및 디지털 전환을 핵심 과제로 추진하고 있음.")

    add_formatted_paragraph(doc, "그러나 현재 한의원 진료 현장에는 다음과 같은 문제가 존재함.")

    p1 = doc.add_paragraph()
    run1 = p1.add_run("첫째, ")
    run1.font.name = '맑은 고딕'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run1_bold = p1.add_run("진료 품질의 편차 문제")
    run1_bold.bold = True
    run1_bold.font.name = '맑은 고딕'
    run1_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run1_cont = p1.add_run("임. 한의학 진료는 개인 경험과 직관에 크게 의존하여, "
        "같은 환자가 다른 한의사에게 진료를 받으면 전혀 다른 처방을 받는 경우가 빈번함.")
    run1_cont.font.name = '맑은 고딕'
    run1_cont._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    p2 = doc.add_paragraph()
    run2 = p2.add_run("둘째, ")
    run2.font.name = '맑은 고딕'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run2_bold = p2.add_run("임상 지식 전수의 단절 문제")
    run2_bold.bold = True
    run2_bold.font.name = '맑은 고딕'
    run2_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run2_cont = p2.add_run("임. 40년 이상 경력을 가진 원로 한의사들의 귀중한 치험례"
        "(실제 치료 경험 기록)가 체계적으로 정리되지 못하고 은퇴와 함께 사라지고 있음.")
    run2_cont.font.name = '맑은 고딕'
    run2_cont._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    p3 = doc.add_paragraph()
    run3 = p3.add_run("셋째, ")
    run3.font.name = '맑은 고딕'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run3_bold = p3.add_run("안전성 확인의 어려움")
    run3_bold.bold = True
    run3_bold.font.name = '맑은 고딕'
    run3_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run3_cont = p3.add_run("임. 양약과 한약을 동시에 복용하는 환자가 많지만, "
        "상호작용(서로 영향을 주고받는 것)을 신속하게 확인할 수 있는 도구가 부족함.")
    run3_cont.font.name = '맑은 고딕'
    run3_cont._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_paragraph()

    # 1-2. 솔루션 개요
    add_heading_custom(doc, '1-2. 솔루션 개요', level=2)

    p_sol = doc.add_paragraph()
    run_sol1 = p_sol.add_run("'온고지신 AI'는 ")
    run_sol1.font.name = '맑은 고딕'
    run_sol1._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run_sol_bold = p_sol.add_run("AI(인공지능) 기반 한의학 CDSS(Clinical Decision Support System, 임상의사결정지원시스템)")
    run_sol_bold.bold = True
    run_sol_bold.font.name = '맑은 고딕'
    run_sol_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run_sol2 = p_sol.add_run("임. 쉽게 말해, 한의사가 환자를 진료할 때 AI가 옆에서 "
        "\"이런 증상에는 이런 처방이 효과가 좋았습니다\"라고 조언해주는 똑똑한 진료 보조 프로그램임.")
    run_sol2.font.name = '맑은 고딕'
    run_sol2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    add_formatted_paragraph(doc, "핵심 기능은 다음과 같음.")

    features = [
        "• 6,000건 이상의 검증된 치험례를 바탕으로 AI가 처방을 추천함",
        "• 환자 증상을 입력하면 유사한 과거 치료 사례를 자동으로 검색함",
        "• 양약과 한약의 상호작용을 자동으로 검사하여 안전한 처방을 지원함"
    ]
    for feature in features:
        add_formatted_paragraph(doc, feature, space_after=3)

    doc.add_paragraph()

    # 1-3. 미래 작용 및 기대 효과
    add_heading_custom(doc, '1-3. 미래 작용 및 기대 효과', level=2)

    add_formatted_paragraph(doc, "본 서비스가 확산되면 다음과 같은 변화가 예상됨.")

    effects = [
        ("진료 표준화 달성", "전국 어디서나 일정 수준 이상의 한의학 진료를 받을 수 있게 됨."),
        ("한의학 지식의 영구 보존", "원로 한의사들의 40년 임상 경험이 디지털 데이터로 전환되어 후대에 전승됨."),
        ("환자 안전성 향상", "양약-한약 상호작용 자동 검사로 복합 처방의 안전성이 크게 높아짐."),
        ("한의학 해외 진출 기반 마련", "체계화된 AI 시스템을 통해 한의학의 과학적 근거를 확보하고, 글로벌 전통의학 시장 진출이 가능해짐.")
    ]

    for i, (title_text, desc) in enumerate(effects, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{['첫째', '둘째', '셋째', '넷째'][i-1]}, ")
        run_num.font.name = '맑은 고딕'
        run_num._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run_bold = p.add_run(title_text)
        run_bold.bold = True
        run_bold.font.name = '맑은 고딕'
        run_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run_desc = p.add_run(f"임. {desc}")
        run_desc.font.name = '맑은 고딕'
        run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_paragraph()

    # ========== 2. 사업의 롤모델 및 차별성 ==========
    add_heading_custom(doc, '2. 사업의 롤모델 및 차별성', level=1)

    # 2-1. 국내외 롤모델 분석
    add_heading_custom(doc, '2-1. 국내외 롤모델 분석', level=2)

    rolemodel_headers = ['구분', '서비스명', '특징', '한계점']
    rolemodel_data = [
        ['국내', '닥터앤서 2.0', '정부 주도 AI 진단 보조, 12개 질환 대상', '양방 중심, 한의학 미지원'],
        ['국내', '루닛', '흉부 X-ray AI 분석, 글로벌 진출 성공', '영상 진단 특화, 처방 추천 없음'],
        ['해외', 'IBM Watson for Oncology', '암 환자 치료 방향 제시', '서양의학 전용, 고가'],
        ['해외', 'Ping An Good Doctor (중국)', 'AI 기반 원격진료, 4억 명 사용자', '중의학 일부 지원, 한의학 미지원']
    ]
    add_table_with_header(doc, rolemodel_headers, rolemodel_data, [2, 4, 5, 5])

    add_formatted_paragraph(doc,
        "위 사례들은 AI가 의료 현장에서 충분히 활용될 수 있음을 증명함. "
        "특히 중국의 'Ping An Good Doctor'는 전통의학과 AI를 결합한 성공 사례로, "
        "본 사업의 방향성이 실현 가능함을 보여줌.")

    # 2-2. 온고지신 AI의 차별성
    add_heading_custom(doc, '2-2. 온고지신 AI의 차별성', level=2)

    diff_headers = ['항목', '기존 서비스', '온고지신 AI']
    diff_data = [
        ['적용 분야', '양방(서양의학) 중심', '한의학 전문 특화'],
        ['데이터 기반', '논문, 가이드라인', '실제 치험례 6,000건+ (국내 유일)'],
        ['처방 추천', '단순 정보 제공', '성공률 기반 AI 맞춤 추천'],
        ['학파 지원', '해당 없음', '고방, 후세방, 사상방 등 학파별 분석'],
        ['상호작용 검사', '양약 간 검사만', '양약-한약 통합 검사']
    ]
    add_table_with_header(doc, diff_headers, diff_data, [3, 5.5, 5.5])

    p_key = doc.add_paragraph()
    run_key_bold = p_key.add_run("핵심 차별점: ")
    run_key_bold.bold = True
    run_key_bold.font.name = '맑은 고딕'
    run_key_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run_key_text = p_key.add_run("40년 경력 원로 한의사의 6,000건 이상 실제 치험례 데이터를 보유한 것은 "
        "국내에서 온고지신 AI가 유일함. 이 데이터는 단순한 교과서 내용이 아닌, "
        "실제 환자에게 적용되어 효과가 검증된 치료 기록임.")
    run_key_text.font.name = '맑은 고딕'
    run_key_text._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_paragraph()

    # ========== 3. 제품/서비스 개발 현황 ==========
    add_heading_custom(doc, '3. 제품/서비스 개발 현황', level=1)

    # 3-1. MVP 개발 완료
    add_heading_custom(doc, '3-1. MVP(최소기능제품) 개발 완료', level=2)

    add_formatted_paragraph(doc, "현재 웹 기반 서비스의 MVP 개발이 완료되어 실제 동작하는 상태임.")

    # 그림 1 - 랜딩 페이지
    add_formatted_paragraph(doc, "[그림 1] 서비스 메인 화면", bold=True)
    add_image_with_caption(doc, os.path.join(SCREENSHOTS_DIR, "01_landing.png"),
        "그림 1. 온고지신 AI 서비스 메인 화면")

    add_formatted_paragraph(doc,
        "서비스 메인 화면에서는 '옛것을 익혀 새것을 안다'는 온고지신의 의미를 담아, "
        "AI가 변증(한의학에서 환자의 상태를 종합적으로 판단하는 것)을 분석해준다는 핵심 가치를 전달함.")

    doc.add_paragraph()

    # 그림 2 - 대시보드
    add_formatted_paragraph(doc, "[그림 2] AI 처방 어시스턴트 대시보드", bold=True)
    add_image_with_caption(doc, os.path.join(SCREENSHOTS_DIR, "03_dashboard.png"),
        "그림 2. AI 처방 어시스턴트 대시보드 화면")

    add_formatted_paragraph(doc, "대시보드에서는 다음 정보를 한눈에 확인할 수 있음.")
    dashboard_info = [
        "• 94.2%: 학습된 치험례의 치료 성공률",
        "• 3분: 평균 AI 분석 소요 시간",
        "• 6,000건+: 학습 완료된 치험례 수"
    ]
    for info in dashboard_info:
        add_formatted_paragraph(doc, info, space_after=3)
    add_formatted_paragraph(doc, "한의사는 이 화면에서 바로 환자 증상을 입력하여 AI 처방 추천을 받을 수 있음.")

    doc.add_paragraph()

    # 그림 3 - AI 진료 어시스턴트
    add_formatted_paragraph(doc, "[그림 3] AI 진료 어시스턴트 - 증상 입력 화면", bold=True)
    add_image_with_caption(doc, os.path.join(SCREENSHOTS_DIR, "05_consultation.png"),
        "그림 3. AI 진료 어시스턴트 증상 입력 화면")

    add_formatted_paragraph(doc,
        "환자 증상 입력 화면에서는 자연어(일상적인 말)로 증상을 입력할 수 있음. "
        "예를 들어 \"65세 남자, 소화가 안되고 배가 차갑습니다. 밥을 먹으면 더부룩하고 설사를 자주 합니다.\"와 "
        "같이 입력하면 됨.")

    add_formatted_paragraph(doc,
        "빠른 입력을 위해 자주 사용되는 증상 버튼(두통, 어지러움, 피로 등)과 "
        "사상체질(태양인, 태음인, 소양인, 소음인) 선택 기능도 제공함.")

    doc.add_paragraph()

    # 그림 4 - 통합 검색
    add_formatted_paragraph(doc, "[그림 4] 통합 검색 기능", bold=True)
    add_image_with_caption(doc, os.path.join(SCREENSHOTS_DIR, "04_unified_search.png"),
        "그림 4. 통합 검색 화면")

    add_formatted_paragraph(doc, "통합 검색 화면에서는 다음 기능을 제공함.")
    search_features = [
        "• 치험례 검색: 6,115건의 실제 치료 기록에서 유사 사례 검색",
        "• 처방 검색: 처방명으로 상세 정보 조회",
        "• 변증 검색: 한의학 진단 유형별 검색",
        "• 학파별 필터: 고방(옛 처방), 후세방(후대 처방), 사상방(체질별 처방) 분류"
    ]
    for feature in search_features:
        add_formatted_paragraph(doc, feature, space_after=3)

    doc.add_paragraph()

    # 그림 5 - 사상체질 진단
    add_formatted_paragraph(doc, "[그림 5] 사상체질 진단 기능", bold=True)
    add_image_with_caption(doc, os.path.join(SCREENSHOTS_DIR, "08_constitution.png"),
        "그림 5. 사상체질 진단 화면")

    add_formatted_paragraph(doc,
        "사상체질(사람을 4가지 체질로 분류하는 한의학 이론) 진단 기능에서는 "
        "12가지 질문을 통해 환자의 체질을 파악함. \"본인의 체형은 어떤 편인가요?\"와 같은 "
        "쉬운 질문으로 구성되어 있어, 환자가 직접 답변할 수도 있음.")

    doc.add_paragraph()

    # 그림 6 - 상호작용 검사
    add_formatted_paragraph(doc, "[그림 6] 양약-한약 상호작용 검사", bold=True)
    add_image_with_caption(doc, os.path.join(SCREENSHOTS_DIR, "09_interactions.png"),
        "그림 6. 양약-한약 상호작용 검사 화면")

    p_inter = doc.add_paragraph()
    run_inter_bold = p_inter.add_run("양약-한약 상호작용 검사")
    run_inter_bold.bold = True
    run_inter_bold.font.name = '맑은 고딕'
    run_inter_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run_inter_text = p_inter.add_run("는 환자 안전을 위한 핵심 기능임. "
        "처방할 한약재와 환자가 복용 중인 양약(예: 와파린, 아스피린)을 입력하면, "
        "두 약물 간의 상호작용 여부를 자동으로 확인해줌.")
    run_inter_text.font.name = '맑은 고딕'
    run_inter_text._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    add_formatted_paragraph(doc,
        "이 기능은 양방과 한방을 병행하는 환자가 많은 현실에서 매우 중요한 안전장치 역할을 함.")

    doc.add_paragraph()

    # 3-2. 기술 스택
    add_heading_custom(doc, '3-2. 기술 스택 및 아키텍처', level=2)

    tech_headers = ['구분', '기술', '선정 이유']
    tech_data = [
        ['AI 엔진', 'Claude API + RAG', '최신 LLM(대규모언어모델)로 자연어 이해 우수'],
        ['프론트엔드', 'React + TypeScript', '안정적인 웹 개발, 유지보수 용이'],
        ['백엔드', 'NestJS + PostgreSQL', '확장성 높은 서버 구조'],
        ['배포', 'Vercel + Fly.io', '클라우드 기반으로 안정적 운영']
    ]
    add_table_with_header(doc, tech_headers, tech_data, [3, 4, 7])

    # ========== 4. 개발 로드맵 ==========
    add_heading_custom(doc, '4. 개발 로드맵', level=1)

    # 4-1. 단계별 개발 계획
    add_heading_custom(doc, '4-1. 단계별 개발 계획', level=2)

    roadmap_headers = ['단계', '시기', '주요 개발 내용', '완료 기준']
    roadmap_data = [
        ['Phase 1', '2026 Q1', 'MVP 고도화, 베타 테스트', '10개 한의원 시범 도입'],
        ['Phase 2', '2026 Q2', '정식 출시, EMR(전자의무기록) 연동', 'EMR 2개사 연동 완료'],
        ['Phase 3', '2026 H2', '모바일 앱 출시, 기능 확장', '100개소 도입'],
        ['Phase 4', '2027', '해외 시장 진출 검토', '500개소 달성']
    ]
    add_table_with_header(doc, roadmap_headers, roadmap_data, [2.5, 2.5, 5, 4])

    # 4-2. 세부 기능 개발 계획
    add_heading_custom(doc, '4-2. 세부 기능 개발 계획', level=2)

    add_formatted_paragraph(doc, "Phase 1 (핵심 기능 완성)", bold=True)
    phase1_items = [
        "• AI 변증 분석 정확도 향상 (목표: 90% 이상)",
        "• 학파별 처방 분류 시스템 구축",
        "• 사용자 피드백 수집 및 반영"
    ]
    for item in phase1_items:
        add_formatted_paragraph(doc, item, space_after=3)

    add_formatted_paragraph(doc, "Phase 2 (시장 확장)", bold=True)
    phase2_items = [
        "• 주요 EMR 업체(한의타임, 오아시스 등)와 연동 개발",
        "• 구독 결제 시스템 고도화",
        "• 고객 지원 체계 구축"
    ]
    for item in phase2_items:
        add_formatted_paragraph(doc, item, space_after=3)

    add_formatted_paragraph(doc, "Phase 3 (서비스 확장)", bold=True)
    phase3_items = [
        "• iOS/Android 모바일 앱 개발",
        "• 음성 입력 기능 추가",
        "• 팔강변증(8가지 기준으로 진단하는 방법) 자동 분석기 개발"
    ]
    for item in phase3_items:
        add_formatted_paragraph(doc, item, space_after=3)

    doc.add_paragraph()

    # ========== 5. 실현 가능성 근거 ==========
    add_heading_custom(doc, '5. 실현 가능성 근거', level=1)

    # 5-1. 기술적 실현 가능성
    add_heading_custom(doc, '5-1. 기술적 실현 가능성', level=2)

    tech_feasibility_headers = ['항목', '현황', '근거']
    tech_feasibility_data = [
        ['AI 기술', '상용화 단계', 'Claude API, GPT-4 등 LLM 기술 성숙'],
        ['데이터', '확보 완료', '6,000건+ 치험례 디지털화 완료'],
        ['개발', 'MVP 완료', '웹 서비스 실제 동작 중'],
        ['인프라', '구축 완료', '클라우드 기반 안정적 운영']
    ]
    add_table_with_header(doc, tech_feasibility_headers, tech_feasibility_data, [3, 3, 8])

    # 5-2. 시장 실현 가능성
    add_heading_custom(doc, '5-2. 시장 실현 가능성', level=2)

    market_items = [
        "• 국내 한의원 수: 14,000개소 이상",
        "• 한의사 수: 25,000명 이상",
        "• 디지털 전환 수요: 코로나19 이후 비대면 진료, AI 도입 관심 급증",
        "• 정부 정책: 한의약육성법, 디지털 헬스케어 육성 정책과 부합"
    ]
    for item in market_items:
        add_formatted_paragraph(doc, item, space_after=3)

    doc.add_paragraph()

    # 5-3. 팀 역량
    add_heading_custom(doc, '5-3. 팀 역량', level=2)

    team_headers = ['역할', '담당자', '주요 역량']
    team_data = [
        ['CEO', '양보름', '사업 기획, 전략 수립, 영업'],
        ['CTO', '이현석', 'AI/ML 개발, 시스템 설계, 풀스택 개발'],
        ['의료자문', '이종대 원장', '한의학 임상 40년, 치험례 데이터 제공']
    ]
    add_table_with_header(doc, team_headers, team_data, [3, 3, 8])

    # ========== 6. 정부 정책과의 연계 ==========
    add_heading_custom(doc, '6. 정부 정책과의 연계', level=1)

    add_formatted_paragraph(doc, "본 사업은 다음 정부 정책과 직접적으로 연계됨.")

    policy_headers = ['정책', '연계 내용']
    policy_data = [
        ['제4차 한의약육성발전종합계획', '한의학 과학화, 디지털 전환 목표와 부합'],
        ['디지털 헬스케어 육성 정책', 'AI 기반 의료 서비스 혁신'],
        ['고령사회 대응 정책', '만성질환 관리, 한방 의료 접근성 향상'],
        ['K-전통의학 세계화', '한의학 표준화를 통한 해외 진출 기반']
    ]
    add_table_with_header(doc, policy_headers, policy_data, [5, 9])

    # ========== 7. 결론 ==========
    add_heading_custom(doc, '7. 결론', level=1)

    add_formatted_paragraph(doc,
        "'온고지신 AI'는 검증된 AI 기술과 국내 유일의 대규모 치험례 데이터를 결합하여, "
        "한의학 진료의 품질 향상과 표준화를 실현할 수 있는 서비스임. "
        "이미 MVP 개발이 완료되어 실제 동작하는 서비스를 보유하고 있으며, "
        "국내외 유사 서비스의 성공 사례를 통해 시장 가능성이 검증됨.")

    add_formatted_paragraph(doc,
        "정부의 한의학 디지털 전환 정책과 맞물려, 본 사업은 한의학의 과학화와 세계화에 기여하는 동시에, "
        "환자에게는 더 안전하고 효과적인 진료를, 한의사에게는 진료 품질 향상을 위한 강력한 도구를 제공할 것임.")

    # 저장
    doc.save(OUTPUT_PATH)
    print(f"문서가 저장되었습니다: {OUTPUT_PATH}")
    return OUTPUT_PATH

if __name__ == "__main__":
    create_document()
