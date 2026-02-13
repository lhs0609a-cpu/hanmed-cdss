#!/usr/bin/env python3
"""
초기창업패키지 사업계획서를 Word 문서(.docx)로 생성하는 스크립트
볼드체가 포함된 전문적인 형식으로 작성
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def set_run_font(run, font_name='맑은 고딕', size=11, bold=False):
    """런(텍스트 조각)의 폰트 설정"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_styled_paragraph(doc, text, style='Normal', font_size=11, bold=False, alignment=None, space_after=12):
    """스타일이 적용된 문단 추가 - **텍스트** 는 볼드 처리"""
    p = doc.add_paragraph()

    if alignment:
        p.alignment = alignment

    p.paragraph_format.space_after = Pt(space_after)

    # **텍스트** 패턴을 찾아 볼드 처리
    pattern = r'\*\*(.*?)\*\*'
    parts = re.split(pattern, text)

    is_bold_part = False
    for i, part in enumerate(parts):
        if not part:
            continue
        # 홀수 인덱스는 볼드 텍스트 (패턴 매치된 그룹)
        is_bold_part = (i % 2 == 1)
        run = p.add_run(part)
        set_run_font(run, size=font_size, bold=bold or is_bold_part)

    return p


def add_section_header(doc, text, level=1):
    """섹션 헤더 추가"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if level == 1:
        # 대분류 헤더
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True)
    elif level == 2:
        # 중분류 헤더
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        set_run_font(run, size=14, bold=True)

    return p


def add_sub_header(doc, text):
    """소분류 헤더 추가"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True)
    return p


def add_bullet_point(doc, text, indent=0.5):
    """글머리 기호 문단 추가"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)

    # **텍스트** 패턴 처리
    pattern = r'\*\*(.*?)\*\*'
    parts = re.split(pattern, text)

    for i, part in enumerate(parts):
        if not part:
            continue
        is_bold_part = (i % 2 == 1)
        run = p.add_run(part)
        set_run_font(run, size=11, bold=is_bold_part)

    return p


def add_separator(doc):
    """구분선 추가"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 40)
    set_run_font(run, size=10)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def create_business_plan():
    """사업계획서 Word 문서 생성"""
    doc = Document()

    # 문서 기본 설정
    sections = doc.sections
    for section in sections:
        section.page_width = Cm(21)  # A4
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ===== 제목 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run('초기창업패키지 사업계획서')
    set_run_font(run, size=24, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    run = subtitle.add_run('온고지신 AI')
    set_run_font(run, size=18, bold=True)

    add_separator(doc)

    # ===== BUMZ =====
    add_section_header(doc, '【 BUMZ (Business Unique Model Zone) 】')

    add_styled_paragraph(doc, '사업 분야: AI 기반 한의학 임상의사결정지원시스템', bold=True)

    add_styled_paragraph(doc,
        '온고지신 AI는 인공지능으로 한의사의 진료를 돕는 소프트웨어 서비스임. '
        '40년 경력 원로 한의사의 실제 치료 사례 6,000건 이상을 학습한 인공지능이 '
        '환자 증상을 분석하고, 적합한 한약 처방과 치료 방향을 추천함.')

    add_styled_paragraph(doc,
        '한의사 개인의 경험과 감에 의존하던 기존 진료 방식에서 벗어나, '
        '**데이터와 인공지능 기반의 객관적이고 체계적인 진료**를 가능하게 함.')

    add_separator(doc)

    # ===== 아이템 개요 =====
    add_section_header(doc, '【 아이템 개요 】')

    add_sub_header(doc, '1. 아이템 탄생 배경')
    add_styled_paragraph(doc,
        '한의학은 수천 년 역사를 가진 대한민국의 소중한 의료 자산임. 그러나 현재 '
        '한의학 지식 대부분이 경험 많은 한의사 머릿속에만 존재하고, 체계적으로 '
        '기록되거나 전수되지 못하고 있음.')
    add_styled_paragraph(doc,
        '40년 이상 경력의 원로 한의사들이 은퇴하면서 수십 년간 쌓은 귀중한 임상 '
        '경험(실제 환자 치료하며 얻은 노하우)이 함께 사라지는 중임. 이 문제를 '
        '해결하기 위해 원로 한의사의 치료 경험을 인공지능에 학습시켜 후배 '
        '한의사들에게 전달하는 시스템을 개발함.')

    add_sub_header(doc, '2. 왜 필요한가')
    add_styled_paragraph(doc, '첫째, **한의사들의 진료 품질 격차를 줄여줌.**')
    add_styled_paragraph(doc,
        '갓 개원한 한의사와 30년 경력 한의사 사이에는 큰 실력 차이가 존재함. '
        '온고지신 AI가 경험 부족한 한의사에게 베테랑 한의사의 지혜를 빌려줘서, '
        '어느 한의원을 가도 양질의 진료를 받을 수 있게 함.')

    add_styled_paragraph(doc, '둘째, **환자 안전을 높여줌.**')
    add_styled_paragraph(doc,
        '양약(병원에서 처방받는 일반 약)과 한약을 함께 복용할 때 생길 수 있는 '
        '부작용을 자동으로 검사해줌. 예를 들어 고혈압 약 먹는 환자에게 처방하면 '
        '안 되는 한약재를 미리 알려줘서 의료 사고를 예방함.')

    add_styled_paragraph(doc, '셋째, **한의학의 과학화와 세계화에 기여함.**')
    add_styled_paragraph(doc,
        '인공지능이 학습한 수천 건의 치료 사례가 한의학 효과를 데이터로 증명하는 '
        '근거가 됨. 한의학이 전 세계에서 인정받는 의학으로 발전하는 데 핵심 역할을 함.')

    add_sub_header(doc, '3. 미래에 어떤 역할을 하는가')
    add_styled_paragraph(doc,
        '가까운 미래에 온고지신 AI는 단순한 진료 보조 도구를 넘어 '
        '**한의학 지식의 종합 플랫폼**으로 발전할 것임.')

    add_bullet_point(doc, '전국 한의사들이 자신의 치료 경험을 공유하고 서로 배우는 공간이 됨')
    add_bullet_point(doc, '환자들이 건강 상태를 기록하고 한의사와 실시간 소통할 수 있음')
    add_bullet_point(doc, '인공지능이 축적된 데이터를 분석해 새로운 치료법 발견에 활용됨')
    add_bullet_point(doc, '해외에서 한의학 진료 원하는 사람들에게 원격 서비스 제공 가능함')

    add_styled_paragraph(doc,
        '궁극적으로 온고지신 AI는 대한민국 한의학의 우수성을 전 세계에 알리고, '
        'K-메디(한국 의료 서비스의 세계 브랜드)의 핵심 축으로 자리 잡을 것임.')

    add_separator(doc)

    # ===== 문제 인식 =====
    add_section_header(doc, '【 문제 인식(Problem)_창업 아이템의 필요성 】')

    add_sub_header(doc, '1. 국가 정책 기반의 문제 인식')

    p = doc.add_paragraph()
    run = p.add_run('이재명 정부 핵심 국정과제와 한의학 현실의 괴리')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '이재명 정부는 "AI 대전환", "지역 의료 살리기", "K-헬스케어 글로벌 진출", '
        '"초고령사회 건강수명 연장"을 핵심 국정과제로 추진 중임. 특히 **1차 의료기관'
        '(동네 병·의원, 한의원)에 AI 기술을 보급**하여 대형병원 쏠림 현상을 해소하고 '
        '전 국민이 양질의 의료 서비스를 받을 수 있게 하는 것이 목표임.')

    add_styled_paragraph(doc,
        '그러나 **한의학 분야는 이 정책을 실현할 기반 자체가 무너지고 있음.** '
        '원로 한의사들의 은퇴로 귀중한 임상 경험이 사라지고, 신규 한의사들은 경험 부족으로 '
        '진료에 어려움을 겪으며, 한의학 데이터의 과학화·표준화는 요원한 상태임.')

    add_sub_header(doc, '2. 구체적 문제 상황')

    # 문제 상황 표
    problem_table = doc.add_table(rows=5, cols=3)
    problem_table.style = 'Table Grid'

    headers = ['문제', '현황', '정부 정책과의 충돌']
    for i, header in enumerate(headers):
        cell = problem_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    problem_data = [
        ['한의학 지식 단절', '30년+ 경력 원로 한의사 15%뿐, 매년 은퇴 가속화', 'AI 학습 데이터 소멸 → AI 대전환 불가'],
        ['신규 한의사 역량 격차', '개원 5년 미만 78%가 진료 자신감 부족', '동네 한의원 경쟁력 약화 → 지역 의료 붕괴'],
        ['환자 안전 시스템 부재', '65세 이상 70%가 양약+한약 복용, 상호작용 검사 無', '노인 환자 안전 위협 → 건강수명 정책 실패'],
        ['한의학 비표준화', '동일 증상에 한의사마다 다른 처방', '과학적 근거 부족 → K-헬스케어 해외 진출 불가'],
    ]

    for row_idx, row_data in enumerate(problem_data):
        for col_idx, cell_data in enumerate(row_data):
            problem_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    add_sub_header(doc, '3. 창업 아이템의 필요성')

    p = doc.add_paragraph()
    run = p.add_run('왜 지금 온고지신 AI가 필요한가?')
    set_run_font(run, size=11, bold=True)

    p = doc.add_paragraph()
    run = p.add_run('첫째, 시간이 없음.')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '40년 경력 원로 한의사들이 매년 은퇴하고 있음. 이들의 머릿속에만 있는 6,000건 이상의 '
        '성공적인 치료 경험이 **기록되지 않으면 영영 사라짐.** 한의학 AI의 학습 데이터가 될 '
        '귀중한 자산이 소멸 위기에 처함. **지금 행동하지 않으면 정부의 한의학 과학화 정책은 '
        '실현 기반 자체를 잃게 됨.**')

    p = doc.add_paragraph()
    run = p.add_run('둘째, 대안이 없음.')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '현재 시장에 한의학 전용 AI 임상의사결정지원시스템(CDSS)은 **단 하나도 없음.** '
        '양방(일반 병원) 분야에는 뷰노, 닥터앤서 등 다수의 AI 솔루션이 존재하지만, 한의학 고유의 '
        '변증(辨證), 사상체질, 한약 처방 체계를 이해하는 AI는 전무함. '
        '**온고지신 AI가 이 공백을 메우는 유일한 솔루션임.**')

    p = doc.add_paragraph()
    run = p.add_run('셋째, 정부 정책과 완벽히 정렬됨.')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc, '온고지신 AI는 이재명 정부 4대 국정과제에 직접 기여함:')

    add_bullet_point(doc, '**AI 대전환**: 한의학 분야 최초 AI CDSS로 정책 선도')
    add_bullet_point(doc, '**지역 의료 강화**: 동네 한의원 진료 품질을 대형 한방병원 수준으로 향상')
    add_bullet_point(doc, '**K-헬스케어 수출**: 한의학 데이터 표준화로 글로벌 진출 기반 구축')
    add_bullet_point(doc, '**초고령사회 대응**: 양약-한약 상호작용 자동 검사로 노인 환자 안전 보장')

    p = doc.add_paragraph()
    run = p.add_run('넷째, 사회적 가치 창출.')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '한의학은 수천 년 역사를 가진 대한민국의 무형 문화자산임. 온고지신 AI는 사라져가는 '
        '한의학 지식을 디지털로 보존하고, 다음 세대 한의사들에게 전수하며, 전 세계에 한의학의 '
        '우수성을 알리는 **문화적·의료적 가치를 동시에 창출**함.')

    add_sub_header(doc, '4. 문제 해결 시 기대 효과')

    # 기대 효과 표
    effect_table = doc.add_table(rows=5, cols=2)
    effect_table.style = 'Table Grid'

    effect_headers = ['이해관계자', '기대 효과']
    for i, header in enumerate(effect_headers):
        cell = effect_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    effect_data = [
        ['한의사', '베테랑 한의사 수준의 진료 역량 확보, 의료사고 위험 감소'],
        ['환자', '어느 한의원을 가도 양질의 진료, 양약-한약 부작용 예방'],
        ['정부', 'AI 대전환·지역 의료·K-헬스케어 정책 목표 달성'],
        ['한의학계', '임상 데이터 축적으로 과학화·세계화 기반 마련'],
    ]

    for row_idx, row_data in enumerate(effect_data):
        for col_idx, cell_data in enumerate(row_data):
            effect_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    add_styled_paragraph(doc,
        '**결론: 온고지신 AI는 "있으면 좋은 서비스"가 아니라, '
        '한의학의 미래를 위해 "반드시 있어야 하는 서비스"임.**')

    add_separator(doc)

    # ===== 정책 연계 상세 =====
    add_section_header(doc, '【 이재명 정부 국정과제 상세 연계 】')

    add_sub_header(doc, '1. 이재명 정부 국정과제와 연결된 핵심 문제')

    p = doc.add_paragraph()
    run = p.add_run('국정과제 1: AI 대전환과 디지털 헬스케어 산업 육성')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '이재명 정부는 **"AI 대전환"**을 국가 핵심 전략으로 추진 중임. 2025년 발표된 '
        '"AI 강국 도약 전략"에 따르면, 의료 분야 AI 적용을 통해 2030년까지 '
        '디지털 헬스케어 시장을 30조 원 규모로 성장시킬 계획임. 특히 **1차 의료기관 '
        '(동네 병·의원, 한의원)에 AI 기술 보급**을 중점 과제로 설정함.')

    p = doc.add_paragraph()
    run = p.add_run('국정과제 2: 의료 공공성 강화와 지역 의료 살리기')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '이재명 정부는 대형병원 쏠림 현상을 해소하고 동네 의원·한의원 경쟁력 강화를 '
        '추진 중임. **AI 진료 지원 시스템을 통해 1차 의료기관의 진료 품질을 높이고**, '
        '환자들이 대형병원 가지 않아도 양질의 치료 받을 수 있게 하는 것이 목표임.')

    p = doc.add_paragraph()
    run = p.add_run('국정과제 3: K-헬스케어 글로벌 진출')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '정부는 K-콘텐츠에 이어 **K-헬스케어를 신성장동력**으로 육성 중임. 한의학은 '
        'K-헬스케어의 핵심 콘텐츠이며, 인공지능 기술과 결합해 세계 시장 진출 '
        '가속화할 계획임. 전통의학의 과학화·표준화가 해외 진출의 필수 조건임.')

    p = doc.add_paragraph()
    run = p.add_run('국정과제 4: 초고령사회 대응과 건강수명 연장')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '대한민국은 2025년 65세 이상 인구 20% 넘는 초고령사회 진입함. 이재명 정부는 '
        '"아프지 않고 오래 사는 사회" 실현 위해 **예방 중심 의료와 만성질환 관리** '
        '강화 추진 중임. 한의학은 예방의학과 만성질환 치료에 강점 있어 핵심 역할 기대됨.')

    add_sub_header(doc, '2. 정책 실현을 가로막는 한의학 분야의 문제')

    add_styled_paragraph(doc,
        '정부의 국정과제 달성을 위해 한의학 분야 혁신이 필수적이나, '
        '현재 다음과 같은 심각한 문제들이 존재함.')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('문제 1: 한의학 지식의 단절 위기')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '현재 활동 중인 한의사 약 25,000명 중 30년 이상 경력의 원로 한의사는 '
        '약 15%에 불과함. 이들이 은퇴하면 수십 년간 쌓은 임상 경험이 함께 사라짐. '
        '**정부가 추진하는 한의학 과학화의 기반이 될 데이터 자체가 소멸 위기에 처함.**')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('문제 2: 1차 의료기관(한의원) 경쟁력 약화')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '신규 개원 한의사들이 경험 부족으로 진료에 어려움 겪음. 설문조사 결과, '
        '개원 5년 미만 한의사의 78%가 "복잡한 환자 만나면 자신감 떨어진다"고 답함. '
        '**동네 한의원 경쟁력 강화라는 정부 목표 달성에 걸림돌이 됨.**')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('문제 3: 환자 안전 시스템 부재')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '65세 이상 인구의 약 70%가 양약과 한약을 함께 복용함. 그러나 양약-한약 '
        '상호작용을 체계적으로 확인할 시스템이 없음. **초고령사회에서 노인 환자 '
        '안전 보장이라는 정부 목표와 정면으로 배치되는 상황임.**')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('문제 4: 한의학 표준화 지연으로 해외 진출 어려움')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '같은 증상에 한의사마다 다른 처방 내리는 경우 많음. 개인 경험에 의존하는 '
        '진료 방식이 원인임. **한의학이 표준화되지 않으면 K-헬스케어 해외 진출이라는 '
        '정부 목표 달성 불가능함.**')

    add_sub_header(doc, '3. 정부 정책과 사업의 연결점')
    add_styled_paragraph(doc,
        '온고지신 AI는 이재명 정부 국정과제 달성에 직접 기여하는 솔루션임.')

    add_bullet_point(doc, '**AI 대전환**: 한의학 분야 최초의 AI 임상의사결정지원시스템')
    add_bullet_point(doc, '**지역 의료 강화**: 동네 한의원 진료 품질을 대형 한방병원 수준으로 향상')
    add_bullet_point(doc, '**K-헬스케어 수출**: 한의학 과학화·표준화로 해외 진출 기반 마련')
    add_bullet_point(doc, '**초고령사회 대응**: 양약-한약 상호작용 검사로 노인 환자 안전 보장')

    add_styled_paragraph(doc,
        '**지금 행동 안 하면 원로 한의사들의 귀중한 임상 경험이 영영 사라지고, '
        '정부의 한의학 발전 정책도 실현 기반을 잃게 됨.**')

    add_separator(doc)

    # ===== 실현 가능성 =====
    add_section_header(doc, '【 실현 가능성(Solution)_창업 아이템의 개발 계획 】')

    add_sub_header(doc, '1. 현재 개발 완료 현황 (MVP 완성)')

    add_styled_paragraph(doc,
        '온고지신 AI는 아이디어 단계가 아닌, **이미 작동하는 완성된 제품**임. '
        'CTO가 단독으로 MVP(최소 기능 제품)를 개발 완료하여 현재 서비스 중임.')

    p = doc.add_paragraph()
    run = p.add_run('개발 완료 항목:')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '웹 서비스 (https://ongojisin.co.kr) → **지금 바로 접속 가능**')
    add_bullet_point(doc, '인공지능 추론 엔진 (Claude API 기반)')
    add_bullet_point(doc, '6,000건 치험례 데이터베이스')
    add_bullet_point(doc, '양약-한약 상호작용 검사 기능')
    add_bullet_point(doc, '모바일 환자용 앱 (React Native)')

    add_styled_paragraph(doc,
        '**→ 이미 만들어진 제품으로 사업 시작하므로 기술 개발 실패 위험이 매우 낮음.**')

    # 이미지 1: 대시보드
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[그림 1] 온고지신 AI 메인 대시보드')
    set_run_font(run, size=10, bold=True)

    try:
        doc.add_picture(r'G:\내 드라이브\developer\hanmed-cdss\docs\screenshots\03_dashboard.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        add_styled_paragraph(doc, '(대시보드 화면 이미지)')

    add_styled_paragraph(doc,
        '▲ 한의사가 로그인하면 보이는 메인 화면. 오늘의 진료 현황, 환자 통계, '
        'AI 추천 활용률 등을 한눈에 확인할 수 있음.')

    add_sub_header(doc, '2. 핵심 기능별 개발 현황')

    p = doc.add_paragraph()
    run = p.add_run('기능 1: AI 진료 상담 (핵심 기능)')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '환자 증상을 입력하면 AI가 6,000건의 치험례를 분석하여 **유사 사례, 추천 처방, '
        '변증 분석**을 제공함. 40년 경력 한의사의 진료 노하우를 AI가 학습하여 '
        '신규 한의사도 베테랑 수준의 진료 지원 받을 수 있음.')

    # 이미지 2: AI 진료 상담
    try:
        doc.add_picture(r'G:\내 드라이브\developer\hanmed-cdss\docs\screenshots\05_consultation.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        add_styled_paragraph(doc, '(AI 진료 상담 화면 이미지)')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[그림 2] AI 진료 상담 화면')
    set_run_font(run, size=10, bold=True)

    p = doc.add_paragraph()
    run = p.add_run('기능 2: 치험례 통합 검색')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '증상, 처방명, 한약재 등 다양한 조건으로 6,000건의 치험례를 검색할 수 있음. '
        '**실제 환자가 호전된 성공 사례**만 엄선하여 수록함.')

    # 이미지 3: 통합 검색
    try:
        doc.add_picture(r'G:\내 드라이브\developer\hanmed-cdss\docs\screenshots\04_unified_search.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        add_styled_paragraph(doc, '(통합 검색 화면 이미지)')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[그림 3] 치험례 통합 검색 화면')
    set_run_font(run, size=10, bold=True)

    p = doc.add_paragraph()
    run = p.add_run('기능 3: 양약-한약 상호작용 검사 (환자 안전)')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '환자가 복용 중인 양약과 처방하려는 한약 사이의 상호작용을 **자동으로 검사**함. '
        '위험한 조합을 미리 경고하여 의료 사고를 예방함. **현재 시중에 이 기능을 제공하는 '
        '한의학 서비스는 온고지신 AI가 유일함.**')

    # 이미지 4: 상호작용 검사
    try:
        doc.add_picture(r'G:\내 드라이브\developer\hanmed-cdss\docs\screenshots\09_interactions.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        add_styled_paragraph(doc, '(상호작용 검사 화면 이미지)')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[그림 4] 양약-한약 상호작용 검사 화면')
    set_run_font(run, size=10, bold=True)

    p = doc.add_paragraph()
    run = p.add_run('기능 4: 환자 관리 시스템')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '환자 정보, 진료 이력, 처방 기록을 체계적으로 관리함. '
        '환자별 치료 경과를 추적하고 AI가 맞춤형 치료 방향을 제안함.')

    # 이미지 5: 환자 관리
    try:
        doc.add_picture(r'G:\내 드라이브\developer\hanmed-cdss\docs\screenshots\07_patients.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        add_styled_paragraph(doc, '(환자 관리 화면 이미지)')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[그림 5] 환자 관리 화면')
    set_run_font(run, size=10, bold=True)

    add_sub_header(doc, '3. 기술 스택 및 아키텍처')

    # 기술 스택 표
    tech_table = doc.add_table(rows=6, cols=3)
    tech_table.style = 'Table Grid'

    tech_headers = ['구분', '기술', '선정 이유']
    for i, header in enumerate(tech_headers):
        cell = tech_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    tech_data = [
        ['프론트엔드', 'React + TypeScript', '안정성, 대규모 커뮤니티, 유지보수 용이'],
        ['백엔드', 'NestJS + PostgreSQL', '엔터프라이즈급 확장성, 타입 안정성'],
        ['AI 엔진', 'Claude API (Anthropic)', '최신 LLM, 한국어 성능 우수, 의료 도메인 강점'],
        ['인프라', 'Vercel + Fly.io', '글로벌 CDN, 자동 스케일링, 비용 효율'],
        ['모바일', 'React Native + Expo', '크로스 플랫폼, 빠른 개발, 네이티브 성능'],
    ]

    for row_idx, row_data in enumerate(tech_data):
        for col_idx, cell_data in enumerate(row_data):
            tech_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    add_sub_header(doc, '4. 향후 개발 로드맵')

    # 로드맵 표
    roadmap_table = doc.add_table(rows=5, cols=3)
    roadmap_table.style = 'Table Grid'

    roadmap_headers = ['단계', '기간', '개발 내용']
    for i, header in enumerate(roadmap_headers):
        cell = roadmap_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    roadmap_data = [
        ['Phase 1', '2026 상반기', 'EMR 연동 API, 사용자 피드백 반영 UI 개선'],
        ['Phase 2', '2026 하반기', 'AI 예후 예측 기능, 진료 성과 대시보드'],
        ['Phase 3', '2027년', '다국어 지원 (영어, 일본어), 해외 진출 준비'],
        ['Phase 4', '2028년', '연구 데이터 플랫폼, B2B API 서비스'],
    ]

    for row_idx, row_data in enumerate(roadmap_data):
        for col_idx, cell_data in enumerate(row_data):
            roadmap_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    add_sub_header(doc, '5. 사업의 롤모델 (검증된 성공 사례)')

    add_styled_paragraph(doc,
        '온고지신 AI는 완전히 새로운 시장을 개척하는 게 아님. 이미 양방(일반 병원) '
        '분야에서 CDSS가 성공적으로 운영되고 있으며, 이를 한의학에 적용하는 것임.')

    # 롤모델 표
    role_table = doc.add_table(rows=4, cols=4)
    role_table.style = 'Table Grid'

    role_headers = ['기업명', '분야', '성과', '시사점']
    for i, header in enumerate(role_headers):
        cell = role_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    role_data = [
        ['뷰노(VUNO)', '의료 AI', '코스닥 상장, 매출 300억', '의료 AI B2B SaaS 모델 검증'],
        ['닥터앤서', '정부 CDSS', '전국 대학병원 도입', '정부 정책 연계 가능성'],
        ['업투데이트', '글로벌 의료정보', '190개국, 200만 구독자', '구독 모델 글로벌 확장성'],
    ]

    for row_idx, row_data in enumerate(role_data):
        for col_idx, cell_data in enumerate(row_data):
            role_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    add_styled_paragraph(doc,
        '**→ 양방에서 검증된 CDSS 모델을 한의학에 최초 적용. 시장 수용성과 사업 모델 모두 이미 검증됨.**')

    add_sub_header(doc, '6. 온고지신 AI의 차별성')

    # 차별성 표
    diff_table = doc.add_table(rows=4, cols=3)
    diff_table.style = 'Table Grid'

    diff_headers = ['차별점', '내용', '경쟁 우위']
    for i, header in enumerate(diff_headers):
        cell = diff_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    diff_data = [
        ['독보적 데이터', '40년 경력 한의사 치험례 6,000건', '돈 주고도 못 사는 자산, 진입장벽'],
        ['한의학 전문 AI', '변증, 사상체질 등 한의학 개념 이해', '일반 AI로 대체 불가'],
        ['양약-한약 상호작용', '국내 유일 자동 검사 기능', '환자 안전, 규제 대응 우위'],
    ]

    for row_idx, row_data in enumerate(diff_data):
        for col_idx, cell_data in enumerate(row_data):
            diff_table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    add_sub_header(doc, '7. 시장 진입 가능성')

    add_styled_paragraph(doc,
        '국내 한의원 약 14,000개소 중 EMR(전자의무기록) 사용하는 곳 약 8,000개소임. '
        '이들은 이미 디지털 서비스에 익숙하므로 온고지신 AI의 **즉시 도입 가능한 잠재 고객**임.')

    add_styled_paragraph(doc,
        '초기 목표: 수도권 한의원 500개소. 월 구독료 10만 원 기준 **연간 매출 60억 원** 규모 사업 가능함.')

    add_styled_paragraph(doc,
        '**결론: 온고지신 AI는 이미 완성된 제품, 검증된 비즈니스 모델, 독보적인 데이터 자산을 '
        '보유하고 있어 실현 가능성이 매우 높음.**')

    add_separator(doc)

    # ===== 성장 전략 =====
    add_section_header(doc, '【 성장 전략 】')

    add_sub_header(doc, '1. 단계별 성장 계획')

    p = doc.add_paragraph()
    run = p.add_run('1단계 (2026년 상반기): 시장 진입기')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '무료 베타 서비스 운영으로 초기 사용자 100명 확보')
    add_bullet_point(doc, '사용자 피드백 반영해 서비스 품질 개선')
    add_bullet_point(doc, '한의과대학 및 한의사 커뮤니티와 협력 관계 구축')

    p = doc.add_paragraph()
    run = p.add_run('2단계 (2026년 하반기): 성장기')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '유료 서비스 정식 출시')
    add_bullet_point(doc, 'EMR 업체와 연동해 한의원에 쉽게 도입')
    add_bullet_point(doc, '목표: 200개 한의원 유료 구독')

    p = doc.add_paragraph()
    run = p.add_run('3단계 (2027년): 확장기')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '전국 500개 한의원 도입 달성')
    add_bullet_point(doc, '환자용 모바일 앱 본격 홍보')
    add_bullet_point(doc, '대형 한방병원 및 한의과대학 도입 추진')

    p = doc.add_paragraph()
    run = p.add_run('4단계 (2028년 이후): 도약기')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '해외 시장 진출 (일본, 중국, 동남아시아)')
    add_bullet_point(doc, '한의학 연구 데이터 플랫폼으로 확장')
    add_bullet_point(doc, '연간 매출 100억 원 달성')

    add_sub_header(doc, '2. 수익 모델')
    add_styled_paragraph(doc, 'B2B SaaS 구독 모델(기업 대상 소프트웨어 월정액 제공 방식) 채택함.')

    p = doc.add_paragraph()
    run = p.add_run('Basic 요금제: 월 99,000원')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '치험례 검색 기능')
    add_bullet_point(doc, '월 100건 인공지능 분석')
    add_bullet_point(doc, '기본 고객 지원')

    p = doc.add_paragraph()
    run = p.add_run('Professional 요금제: 월 199,000원')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '무제한 인공지능 분석')
    add_bullet_point(doc, 'EMR 연동')
    add_bullet_point(doc, '양약-한약 상호작용 검사')
    add_bullet_point(doc, '우선 고객 지원')

    p = doc.add_paragraph()
    run = p.add_run('Enterprise 요금제: 별도 협의')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '다지점 한의원 및 한방병원용')
    add_bullet_point(doc, '전용 서버 제공')
    add_bullet_point(doc, '맞춤형 기능 개발')

    add_sub_header(doc, '3. 마케팅 전략')

    p = doc.add_paragraph()
    run = p.add_run('전략 1: 학술 마케팅')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '한의학 관련 학회, 세미나에서 온고지신 AI 소개함. 한의사들은 학술적 근거 '
        '있는 제품 선호하므로, 서비스 효과를 연구 논문으로 발표해 신뢰 확보함.')

    p = doc.add_paragraph()
    run = p.add_run('전략 2: 입소문 마케팅')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '초기 사용자인 한의사들이 동료에게 추천하도록 함. 추천인과 신규 가입자 '
        '모두에게 1개월 무료 사용 혜택 제공함.')

    p = doc.add_paragraph()
    run = p.add_run('전략 3: 교육 기관 협력')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '한의과대학 재학생들에게 무료 서비스 제공해서, 졸업 후 개원할 때 '
        '자연스럽게 유료 고객이 되도록 함.')

    p = doc.add_paragraph()
    run = p.add_run('전략 4: EMR 업체 제휴')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '한의원에서 사용하는 EMR 소프트웨어에 온고지신 AI 기본 탑재해서, '
        '한의사들이 별도 설치 없이 바로 사용할 수 있게 함.')

    add_sub_header(doc, '4. 예상 재무 성과')

    # 표 생성
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    headers = ['', '2026년', '2027년', '2028년']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    data = [
        ['유료 구독 한의원', '200개소', '500개소', '1,000개소'],
        ['연간 매출', '24억 원', '60억 원', '120억 원'],
        ['영업이익률', '-20%', '15%', '30%'],
    ]

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()
    add_styled_paragraph(doc,
        '**손익분기점(BEP, 수입과 지출이 같아지는 시점)은 유료 구독 200개소 '
        '달성 시점인 2026년 말로 예상됨.**')

    add_separator(doc)

    # ===== 팀 구성 =====
    add_section_header(doc, '【 팀 구성 】')

    add_sub_header(doc, '1. 핵심 팀원')

    p = doc.add_paragraph()
    run = p.add_run('대표이사 (CEO): 양보름')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '역할: 사업 총괄, 전략 수립, 투자 유치, 영업 총괄')
    add_bullet_point(doc, '강점: 헬스케어 스타트업 경험, 한의사 네트워크 보유')

    p = doc.add_paragraph()
    run = p.add_run('기술이사 (CTO): 이현석')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '역할: 기술 개발 총괄, 인공지능 설계, 서비스 운영')
    add_bullet_point(doc, '강점: 풀스택 개발 역량, 인공지능 서비스 개발 경험')
    add_bullet_point(doc, '개발 실적: **온고지신 AI MVP 단독 개발 완료**')

    p = doc.add_paragraph()
    run = p.add_run('의료자문(1): 안경모 대표원장')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '현직: 소잠한의원 대표원장')
    add_bullet_point(doc, '학력: 대구한의대학교 한의학과 졸업, 대전대학교 한의학과 대학원 졸업')
    add_bullet_point(doc, '면허: 한의사 면허 취득(1994년, **30년 이상 경력**), 미국 한의사 면허(NCCAOM)')
    add_bullet_point(doc, '주요 경력: 대한통합방제한의학회 부회장, 식품의약품안전처 중앙약사심의위원회 전문가')
    add_bullet_point(doc, '기타 활동: 상태의학연구소 고령자채록사업단 단장, 동의한약분석센터 관능검사위원')
    add_bullet_point(doc, '역할: **한의학 총괄 자문, 데이터 품질 검증, 학회 네트워크 연계**')

    p = doc.add_paragraph()
    run = p.add_run('의료자문(2): 이종대 원장')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '역할: 한의학 자문, 치험례 데이터 제공 및 검증')
    add_bullet_point(doc, '경력: 한의원 원장 40년, 치험례 6,000건 이상 보유')
    add_bullet_point(doc, '강점: **국내 최고 수준의 임상 경험과 데이터 보유**')

    add_sub_header(doc, '2. 팀의 강점')

    p = doc.add_paragraph()
    run = p.add_run('강점 1: 기술과 의료의 완벽한 조합')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '기술 담당 CTO와 **30년, 40년 경력의 두 원로 한의사**가 긴밀히 협력해서, '
        '실제 한의사들이 필요로 하는 기능을 정확하게 구현함. 특히 안경모 원장은 '
        '학회 네트워크와 정부 기관 자문 경험 보유해 사업 확장에 핵심 역할 수행함.')

    p = doc.add_paragraph()
    run = p.add_run('강점 2: 이미 작동하는 제품 보유')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '많은 스타트업이 아이디어 단계에서 시작하지만, 온고지신 AI는 이미 개발 '
        '완료된 제품 가지고 있음. CTO가 단독으로 MVP 완성한 것은 뛰어난 기술력의 증거임.')

    p = doc.add_paragraph()
    run = p.add_run('강점 3: 희소한 데이터 자산과 학회 네트워크')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '40년 경력 원로 한의사의 치험례 6,000건은 쉽게 구할 수 없는 귀중한 자산임. '
        '또한 안경모 원장의 **대한통합방제한의학회 부회장, 식약처 전문가 위원** 등 '
        '공식 네트워크를 통해 학계와 정부 기관 협력 기반 확보함. 경쟁사가 쉽게 따라올 수 없는 진입장벽임.')

    add_sub_header(doc, '3. 향후 채용 계획')

    p = doc.add_paragraph()
    run = p.add_run('2026년 상반기:')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '백엔드 개발자 1명 (서버 개발 담당)')
    add_bullet_point(doc, '프론트엔드 개발자 1명 (사용자 화면 개발 담당)')

    p = doc.add_paragraph()
    run = p.add_run('2026년 하반기:')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '마케팅 담당자 1명')
    add_bullet_point(doc, '고객 지원 담당자 1명')

    p = doc.add_paragraph()
    run = p.add_run('2027년:')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '인공지능 엔지니어 1명')
    add_bullet_point(doc, '영업 담당자 2명')

    add_separator(doc)

    # ===== 정부 정책 연계 =====
    add_section_header(doc, '【 이재명 정부 정책과의 연계 】')

    add_styled_paragraph(doc,
        '온고지신 AI는 이재명 정부 핵심 국정과제와 직접 연결되어 정책 실현에 기여함.')

    p = doc.add_paragraph()
    run = p.add_run('1. AI 대전환 정책 (과기정통부·산업부 공동 추진)')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '이재명 정부는 "AI가 바꾸는 대한민국"을 비전으로 전 산업에 AI 기술 확산 추진 중임. '
        '**의료 분야 AI 적용을 최우선 과제로 선정**하고 1차 의료기관 AI 보급 사업 진행 중임. '
        '온고지신 AI는 한의학 분야 AI 적용의 선도 사례로 정책 목표에 직접 부합함.')

    p = doc.add_paragraph()
    run = p.add_run('2. 지역 의료 살리기 정책 (보건복지부)')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '대형병원 쏠림 현상 해소 위해 **동네 의원·한의원 경쟁력 강화** 추진 중임. '
        'AI 진료 지원 시스템 도입 한의원에 건강보험 가산점 부여하는 방안 검토 중임. '
        '온고지신 AI 도입 한의원은 정책 혜택 우선 대상이 될 수 있음.')

    p = doc.add_paragraph()
    run = p.add_run('3. K-헬스케어 글로벌 진출 정책 (복지부·문체부 협력)')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '정부는 K-콘텐츠 성공에 이어 **K-헬스케어를 차세대 수출 주력 산업**으로 육성 중임. '
        '한의학은 K-헬스케어 핵심 콘텐츠이며, 과학화·표준화가 해외 진출 필수 조건임. '
        '온고지신 AI는 한의학 데이터 표준화를 통해 글로벌 진출 기반 마련함.')

    p = doc.add_paragraph()
    run = p.add_run('4. 초고령사회 대응 정책 (범정부 TF)')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '2025년 초고령사회 진입에 따라 **"건강수명 80세 시대"** 목표 설정함. '
        '예방 중심 의료와 만성질환 관리 강화가 핵심 과제임. 한의학이 강점 가진 분야이며, '
        '온고지신 AI는 양약-한약 상호작용 검사로 노인 환자 안전 관리에 직접 기여함.')

    p = doc.add_paragraph()
    run = p.add_run('5. 창업 생태계 활성화 정책 (중기부)')
    set_run_font(run, size=11, bold=True)

    add_styled_paragraph(doc,
        '이재명 정부는 **창업 지원 예산을 역대 최대 규모**로 편성하고, 특히 AI·바이오 '
        '분야 스타트업에 집중 지원 중임. 초기창업패키지 등 정부 지원사업과 연계해 '
        '온고지신 AI의 빠른 성장 기반 확보 가능함.')

    add_separator(doc)

    # ===== 핵심 요약 =====
    add_section_header(doc, '【 핵심 요약 】')

    add_styled_paragraph(doc, '온고지신 AI는 인공지능으로 한의사 진료 돕는 서비스임.')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('핵심 가치:')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '40년 경력 원로 한의사 지혜를 인공지능에 담아 모든 한의사에게 전달')
    add_bullet_point(doc, '양약과 한약 상호작용 자동 검사로 환자 안전 보장')
    add_bullet_point(doc, '한의학 지식의 과학화와 세계화에 기여')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('실현 가능성:')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '양방 분야 성공적인 CDSS 사례를 한의학에 적용 (검증된 모델)')
    add_bullet_point(doc, '이미 개발 완료된 제품과 6,000건 독보적 데이터 보유 (기술 위험 없음)')
    add_bullet_point(doc, '국내 최초 한의학 특화 인공지능으로 시장 선점 (경쟁자 없음)')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('성장 전략:')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '2026년 200개 한의원, 2027년 500개 한의원 도입 목표')
    add_bullet_point(doc, 'B2B SaaS 구독 모델로 안정적 수익 구조 확보')
    add_bullet_point(doc, '2028년 이후 해외 시장 진출')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('팀 역량:')
    set_run_font(run, size=11, bold=True)

    add_bullet_point(doc, '기술, 사업, 의료 아우르는 균형 잡힌 팀 구성')
    add_bullet_point(doc, 'CTO의 MVP 단독 개발로 입증된 기술력')
    add_bullet_point(doc, '30년, 40년 경력 두 원로 한의사 참여로 확보한 신뢰성과 학회 네트워크')

    add_styled_paragraph(doc,
        '온고지신 AI는 사라져가는 한의학의 지혜를 디지털로 보존하고, '
        '누구나 양질의 한의학 진료 받을 수 있는 세상을 만들어 감.')

    add_separator(doc)

    # ===== 문의처 =====
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(24)
    run = footer.add_run('문 의 처')
    set_run_font(run, size=14, bold=True)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('대표이사: 양보름\n기술총괄: 이현석\n서비스: https://hanmed-cdss.vercel.app')
    set_run_font(run, size=11)

    # 저장
    output_path = r'G:\내 드라이브\developer\hanmed-cdss\docs\초기창업패키지_사업계획서_온고지신AI.docx'
    doc.save(output_path)
    print(f'Word 문서가 생성되었습니다: {output_path}')
    return output_path


if __name__ == '__main__':
    create_business_plan()
