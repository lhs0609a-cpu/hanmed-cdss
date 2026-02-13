"""
성장전략(Scale-up)_사업화 추진 전략
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = r"G:\내 드라이브\developer\hanmed-cdss"
SCREENSHOTS_DIR = os.path.join(BASE_DIR, "docs", "screenshots")
OUTPUT_PATH = os.path.join(BASE_DIR, "docs", "성장전략_사업화추진전략.docx")

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_paragraph(doc, text, bold=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_mixed_paragraph(doc, parts):
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return heading

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], 'E3F2FD')

    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    run.font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def add_image(doc, image_path, caption, width=5.0):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))

        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run(caption)
        caption_run.font.size = Pt(9)
        caption_run.font.name = '맑은 고딕'
        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        caption_run.italic = True
        caption_p.paragraph_format.space_after = Pt(10)

def create_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== 제목 ==========
    title = doc.add_heading('성장전략(Scale-up)_사업화 추진 전략', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== 1. 목표시장 분석 ==========
    add_heading(doc, '1. 목표시장 분석', level=1)

    add_heading(doc, '1-1. 시장 규모 (TAM-SAM-SOM)', level=2)

    add_paragraph(doc,
        "TAM(전체시장), SAM(유효시장), SOM(수익시장)을 단계별로 분석하여 현실적인 목표 시장을 설정함.")

    market_headers = ['구분', '정의', '규모', '산출 근거']
    market_data = [
        ['TAM\n(전체시장)', '글로벌 CDSS + 전통의학 시장', '약 100조원', 'CDSS 80억$ + 전통의학 3,500억$'],
        ['SAM\n(유효시장)', '국내 한의원 시장', '약 1,400억원', '한의원 14,000개소 × 연 1,000만원'],
        ['SOM\n(수익시장)', '초기 목표 시장 (수도권)', '약 60억원', '500개소 × 월 10만원 × 12개월']
    ]
    add_table(doc, market_headers, market_data, [2, 4, 3, 5])

    add_mixed_paragraph(doc, [
        ("초기 3년간 ", False),
        ("SOM 60억원 시장의 10% 점유(연 6억원 매출)", True),
        ("를 목표로 함.", False)
    ])

    doc.add_paragraph()

    # 1-2. 목표 고객
    add_heading(doc, '1-2. 목표 고객', level=2)

    customer_headers = ['구분', '특징', '니즈', '규모']
    customer_data = [
        ['1차 타겟\n(신규 한의사)', '경력 5년 미만', '임상 경험 부족 보완, 처방 자신감 필요', '약 5,000명'],
        ['2차 타겟\n(개원 한의사)', '1인 한의원 운영', '진료 효율화, 경쟁력 강화', '약 10,000개소'],
        ['3차 타겟\n(한방병원)', '다수 한의사 근무', '진료 표준화, 교육 도구 필요', '약 300개소']
    ]
    add_table(doc, customer_headers, customer_data, [2.5, 3, 5, 3.5])

    doc.add_paragraph()

    # ========== 2. 비즈니스 모델 ==========
    add_heading(doc, '2. 비즈니스 모델', level=1)

    add_heading(doc, '2-1. 수익 모델: B2B SaaS 구독', level=2)

    add_mixed_paragraph(doc, [
        ("본 서비스는 ", False),
        ("B2B SaaS(기업 대상 구독형 소프트웨어) 모델", True),
        ("로 운영함. 한의원이 월 정액을 내고 서비스를 이용하는 방식임.", False)
    ])

    pricing_headers = ['요금제', '월 요금', '주요 기능', '타겟 고객']
    pricing_data = [
        ['Basic', '99,000원', '기본 검색 + 월 100건 AI 분석', '신규 한의사, 소규모 한의원'],
        ['Professional', '199,000원', '무제한 AI 분석 + EMR 연동', '중견 한의원'],
        ['Enterprise', '별도 협의', '전용 서버 + 맞춤 개발 + 다지점 지원', '한방병원, 프랜차이즈']
    ]
    add_table(doc, pricing_headers, pricing_data, [2.5, 2.5, 5, 4])

    add_image(doc, os.path.join(SCREENSHOTS_DIR, "01_landing.png"),
        "[그림 1] 온고지신 AI 서비스 메인 화면")

    # 2-2. 단위 경제학
    add_heading(doc, '2-2. 단위 경제학 (Unit Economics)', level=2)

    add_paragraph(doc, "고객 1명당 수익성을 분석한 결과, 건전한 비즈니스 구조를 갖추고 있음.")

    unit_headers = ['지표', '수치', '설명']
    unit_data = [
        ['CAC (고객획득비용)', '50만원', '마케팅, 영업 비용을 신규 고객 수로 나눈 값'],
        ['LTV (고객생애가치)', '480만원', '월 10만원 × 48개월 (평균 이용 기간)'],
        ['LTV/CAC', '9.6배', '3배 이상이면 건전한 구조 (업계 기준)'],
        ['BEP (손익분기점)', '200개소', '월 고정비용 대비 필요 고객 수']
    ]
    add_table(doc, unit_headers, unit_data, [3.5, 2.5, 8])

    doc.add_paragraph()

    # ========== 3. 마케팅 및 영업 전략 ==========
    add_heading(doc, '3. 마케팅 및 영업 전략', level=1)

    add_heading(doc, '3-1. 고객 획득 전략 (GTM 전략)', level=2)

    add_paragraph(doc, "GTM(Go-To-Market)은 제품을 시장에 출시하고 고객을 확보하는 전략을 말함.")

    gtm_headers = ['단계', '전략', '세부 내용', '목표']
    gtm_data = [
        ['1단계\n(인지)', '콘텐츠 마케팅', '한의학 관련 블로그, 유튜브 운영\nAI 활용 사례 공유', '월 방문자 1만 명'],
        ['2단계\n(관심)', '무료 체험', '14일 무료 체험 제공\n핵심 기능 경험 유도', '전환율 20%'],
        ['3단계\n(전환)', '학회/세미나', '대한한의사협회 세미나 참여\n한의과대학 특강', '월 50건 문의'],
        ['4단계\n(유지)', '고객 성공팀', '전담 CS 배치\n정기 사용 리포트 제공', '이탈률 5% 미만']
    ]
    add_table(doc, gtm_headers, gtm_data, [2, 3, 6, 3])

    doc.add_paragraph()

    # 3-2. 채널 전략
    add_heading(doc, '3-2. 채널 전략', level=2)

    channel_headers = ['채널', '방법', '예상 효과']
    channel_data = [
        ['온라인 직접 판매', '자사 웹사이트를 통한 가입/결제', '마진율 높음, 고객 데이터 확보'],
        ['EMR 연동 파트너십', '한의타임, 오아시스 등 EMR 업체와 제휴', '기존 고객 대상 접근 용이'],
        ['학회/협회 채널', '대한한의사협회, 각종 학회 공식 추천', '신뢰도 확보, 대량 도입 가능'],
        ['한의과대학 채널', '학생 대상 무료 제공 → 졸업 후 유료 전환', '장기 고객 확보']
    ]
    add_table(doc, channel_headers, channel_data, [3.5, 5.5, 5])

    doc.add_paragraph()

    # ========== 4. 성장 로드맵 ==========
    add_heading(doc, '4. 성장 로드맵', level=1)

    add_heading(doc, '4-1. 연도별 성장 목표', level=2)

    growth_headers = ['연도', '고객 수', '매출', '주요 마일스톤']
    growth_data = [
        ['2026', '100개소', '1.2억원', 'MVP 고도화, 정식 출시, EMR 연동'],
        ['2027', '500개소', '6억원', '모바일 앱, 한방병원 진입'],
        ['2028', '1,500개소', '18억원', '시장 점유율 10%, 시리즈A 투자 유치'],
        ['2029', '3,000개소', '36억원', '해외 시장 진출 (일본, 중국)']
    ]
    add_table(doc, growth_headers, growth_data, [2, 2.5, 2.5, 7])

    add_image(doc, os.path.join(SCREENSHOTS_DIR, "04_unified_search.png"),
        "[그림 2] 치험례 통합 검색 화면 (핵심 경쟁력)")

    # 4-2. 확장 전략
    add_heading(doc, '4-2. 확장 전략', level=2)

    add_paragraph(doc, "단계 1: 수직 확장 (한의학 내 깊이 확대)", bold=True)
    vertical = [
        "• 학파별 전문 모듈 추가 (사상의학, 동의보감 등)",
        "• 침구 처방 추천 기능 확장",
        "• 한약재 재고 관리 연동"
    ]
    for item in vertical:
        add_paragraph(doc, item, space_after=3)

    add_paragraph(doc, "단계 2: 수평 확장 (인접 시장 진출)", bold=True)
    horizontal = [
        "• 한의과대학 교육용 플랫폼 제공",
        "• 일반인 대상 건강 상담 서비스 (B2C)",
        "• 제약회사 대상 한약재 데이터 분석 서비스"
    ]
    for item in horizontal:
        add_paragraph(doc, item, space_after=3)

    add_paragraph(doc, "단계 3: 글로벌 확장", bold=True)
    global_exp = [
        "• 일본 시장: 캄포(漢方) 의학 시장 진출",
        "• 중국 시장: 중의학 CDSS 현지화",
        "• 동남아 시장: 전통의학 수요 증가 지역 공략"
    ]
    for item in global_exp:
        add_paragraph(doc, item, space_after=3)

    doc.add_paragraph()

    # ========== 5. 리스크 관리 ==========
    add_heading(doc, '5. 리스크 관리', level=1)

    risk_headers = ['리스크', '발생 가능성', '대응 전략']
    risk_data = [
        ['경쟁사 출현', '중간', '선점 효과 극대화, 데이터 우위 유지, 특허 출원'],
        ['고객 이탈', '낮음', '고객 성공팀 운영, 지속적 기능 업데이트'],
        ['기술 변화', '중간', 'AI 모델 다변화 (Claude, GPT 등 복수 지원)'],
        ['규제 리스크', '낮음', '의료기기 인허가 사전 검토, 법률 자문 확보']
    ]
    add_table(doc, risk_headers, risk_data, [3, 3, 8])

    doc.add_paragraph()

    # ========== 6. 투자 유치 계획 ==========
    add_heading(doc, '6. 투자 유치 계획', level=1)

    invest_headers = ['단계', '시기', '목표 금액', '용도']
    invest_data = [
        ['Seed', '2026 상반기', '2억원', 'MVP 고도화, 초기 마케팅'],
        ['Pre-A', '2026 하반기', '5억원', 'EMR 연동, 팀 확장'],
        ['Series A', '2028', '20억원', '시장 확대, 해외 진출 준비']
    ]
    add_table(doc, invest_headers, invest_data, [2.5, 3, 3, 5.5])

    add_mixed_paragraph(doc, [
        ("초기창업패키지 선정 시, ", False),
        ("정부 지원금을 Seed 투자금으로 활용", True),
        ("하여 제품 완성도를 높이고, 이를 기반으로 후속 투자를 유치할 계획임.", False)
    ])

    # 저장
    doc.save(OUTPUT_PATH)
    print(f"문서가 저장되었습니다: {OUTPUT_PATH}")
    return OUTPUT_PATH

if __name__ == "__main__":
    create_document()
