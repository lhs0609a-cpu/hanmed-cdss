"""
실현 가능성(Solution) - 창업 아이템의 개발 계획 (간결 버전)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 경로 설정
BASE_DIR = r"G:\내 드라이브\developer\hanmed-cdss"
SCREENSHOTS_DIR = os.path.join(BASE_DIR, "docs", "screenshots")
OUTPUT_PATH = os.path.join(BASE_DIR, "docs", "실현가능성_개발계획.docx")

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_paragraph(doc, text, bold=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_mixed_paragraph(doc, parts):
    """부분적으로 볼드 처리된 문단 추가. parts = [(text, bold), ...]"""
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return heading

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], 'E8F5E9')

    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    run.font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def add_image(doc, image_path, caption, width=5.0):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))

        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run(caption)
        caption_run.font.size = Pt(9)
        caption_run.font.name = '맑은 고딕'
        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        caption_run.italic = True
        caption_p.paragraph_format.space_after = Pt(10)

def create_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== 제목 ==========
    title = doc.add_heading('실현 가능성(Solution)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('창업 아이템의 개발 계획')
    run.font.size = Pt(14)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_paragraph()

    # ========== 1. 아이템 개요 ==========
    add_heading(doc, '1. 아이템 개요', level=1)

    add_mixed_paragraph(doc, [
        ("'온고지신 AI'는 ", False),
        ("AI 기반 한의학 임상의사결정지원시스템(CDSS)", True),
        ("임. CDSS란 의사가 환자를 진료할 때 컴퓨터가 최적의 치료 방법을 추천해주는 시스템을 말함.", False)
    ])

    add_paragraph(doc,
        "쉽게 말해, 한의사가 환자 증상을 입력하면 AI가 \"비슷한 환자에게 이 처방이 효과가 좋았습니다\"라고 "
        "알려주는 똑똑한 진료 보조 프로그램임.")

    doc.add_paragraph()

    # 왜 필요한가?
    add_heading(doc, '▶ 왜 필요한가?', level=2)

    add_paragraph(doc,
        "우리나라는 고령화로 만성질환자가 급증하고, 양방과 한방을 함께 이용하는 환자가 늘어나고 있음. "
        "그러나 한의학 진료 현장에는 해결해야 할 문제가 있음.")

    problems = [
        ("진료 품질 편차", "한의사마다 경험이 달라 같은 환자도 다른 처방을 받는 경우가 많음"),
        ("임상 지식 단절", "원로 한의사의 40년 치료 경험이 체계적으로 전수되지 못하고 사라지고 있음"),
        ("안전성 확인 어려움", "양약과 한약을 함께 복용할 때 부작용을 빠르게 확인할 도구가 부족함")
    ]

    for title_text, desc in problems:
        add_mixed_paragraph(doc, [
            (f"• {title_text}: ", True),
            (desc, False)
        ])

    doc.add_paragraph()

    # 어떻게 해결하는가?
    add_heading(doc, '▶ 어떻게 해결하는가?', level=2)

    add_mixed_paragraph(doc, [
        ("온고지신 AI는 ", False),
        ("40년 경력 원로 한의사의 6,000건 이상 실제 치험례(치료 경험 기록)", True),
        ("를 AI가 학습하여 처방을 추천함. 이를 통해 신규 한의사도 베테랑 수준의 진료가 가능해짐.", False)
    ])

    doc.add_paragraph()

    # 핵심 기능
    add_heading(doc, '▶ 핵심 기능', level=2)

    add_image(doc, os.path.join(SCREENSHOTS_DIR, "03_dashboard.png"),
        "[그림 1] AI 처방 어시스턴트 대시보드")

    features = [
        "① AI 변증 분석: 환자 증상을 입력하면 AI가 한의학적 진단(변증)을 자동 분석함",
        "② 치험례 검색: 6,000건 이상의 실제 치료 기록에서 유사 사례를 즉시 검색함",
        "③ 처방 추천: 성공률이 높은 처방을 우선 추천하여 치료 효과를 높임",
        "④ 양약-한약 상호작용 검사: 환자가 복용 중인 양약과 한약의 부작용 여부를 자동 확인함"
    ]
    for f in features:
        add_paragraph(doc, f, space_after=3)

    doc.add_paragraph()

    # 미래에 어떤 역할을 하는가?
    add_heading(doc, '▶ 미래에 어떤 역할을 하는가?', level=2)

    future = [
        "• 전국 어디서나 균일한 수준의 한의학 진료가 가능해짐",
        "• 원로 한의사의 임상 경험이 디지털로 영구 보존되어 후대에 전승됨",
        "• 한의학의 과학적 근거가 축적되어 해외 진출의 발판이 마련됨"
    ]
    for f in future:
        add_paragraph(doc, f, space_after=3)

    doc.add_paragraph()

    # ========== 2. 롤모델 및 차별성 ==========
    add_heading(doc, '2. 사업의 롤모델 및 차별성', level=1)

    add_heading(doc, '▶ 국내외 롤모델', level=2)

    add_paragraph(doc, "AI 의료 서비스는 이미 국내외에서 성공적으로 운영되고 있음.")

    rolemodel_headers = ['서비스명', '특징', '한계점']
    rolemodel_data = [
        ['닥터앤서 2.0 (국내)', '정부 주도 AI 진단 보조', '양방 중심, 한의학 미지원'],
        ['루닛 (국내)', 'X-ray AI 분석, 글로벌 진출', '영상 진단만 가능'],
        ['Ping An Good Doctor (중국)', 'AI 원격진료, 4억 명 사용', '한의학 미지원']
    ]
    add_table(doc, rolemodel_headers, rolemodel_data, [4, 5, 5])

    add_paragraph(doc,
        "특히 중국의 'Ping An Good Doctor'는 전통의학과 AI를 결합해 성공한 사례로, "
        "본 사업의 방향성이 실현 가능함을 증명함.")

    doc.add_paragraph()

    add_heading(doc, '▶ 온고지신 AI의 차별성', level=2)

    diff_headers = ['항목', '기존 서비스', '온고지신 AI']
    diff_data = [
        ['대상 분야', '양방(서양의학)', '한의학 전문'],
        ['데이터', '논문, 가이드라인', '실제 치험례 6,000건+ (국내 유일)'],
        ['처방 추천', '정보 제공', '성공률 기반 맞춤 추천'],
        ['상호작용', '양약 간만 검사', '양약-한약 통합 검사']
    ]
    add_table(doc, diff_headers, diff_data, [3, 5, 5])

    add_mixed_paragraph(doc, [
        ("핵심 차별점: ", True),
        ("40년 경력 원로 한의사의 실제 치험례 데이터를 보유한 것은 국내에서 온고지신 AI가 유일함.", False)
    ])

    doc.add_paragraph()

    # ========== 3. 개발 현황 ==========
    add_heading(doc, '3. 개발 현황', level=1)

    add_mixed_paragraph(doc, [
        ("현재 ", False),
        ("MVP(최소기능제품) 개발이 완료", True),
        ("되어 실제 동작하는 웹 서비스를 보유하고 있음.", False)
    ])

    doc.add_paragraph()

    # 이미지들
    add_image(doc, os.path.join(SCREENSHOTS_DIR, "05_consultation.png"),
        "[그림 2] AI 진료 어시스턴트 - 증상 입력 화면")

    add_paragraph(doc,
        "환자 증상을 자연어(일상적인 말)로 입력하면 AI가 분석하여 적합한 처방을 추천함. "
        "\"65세 남자, 소화가 안되고 배가 차갑습니다\"와 같이 입력하면 됨.")

    doc.add_paragraph()

    add_image(doc, os.path.join(SCREENSHOTS_DIR, "09_interactions.png"),
        "[그림 3] 양약-한약 상호작용 검사 화면")

    add_paragraph(doc,
        "환자가 복용 중인 양약(예: 아스피린)과 처방할 한약재를 입력하면, "
        "두 약물 간 부작용 여부를 자동으로 확인해줌.")

    doc.add_paragraph()

    # ========== 4. 개발 로드맵 ==========
    add_heading(doc, '4. 개발 로드맵', level=1)

    roadmap_headers = ['단계', '시기', '주요 내용']
    roadmap_data = [
        ['Phase 1', '2026 Q1', '베타 테스트, 10개 한의원 시범 도입'],
        ['Phase 2', '2026 Q2', '정식 출시, EMR(전자의무기록) 연동'],
        ['Phase 3', '2026 하반기', '모바일 앱 출시, 100개소 도입'],
        ['Phase 4', '2027', '500개소 달성, 해외 진출 검토']
    ]
    add_table(doc, roadmap_headers, roadmap_data, [2.5, 2.5, 9])

    doc.add_paragraph()

    # ========== 5. 정부 정책 연계 ==========
    add_heading(doc, '5. 정부 정책과의 연계', level=1)

    add_paragraph(doc, "본 사업은 정부의 핵심 정책과 직접 연계됨.")

    policy_data = [
        ['제4차 한의약육성발전종합계획', '한의학 과학화, 디지털 전환 목표와 부합'],
        ['디지털 헬스케어 육성 정책', 'AI 기반 의료 서비스 혁신'],
        ['고령사회 대응 정책', '만성질환 관리, 한방 의료 접근성 향상']
    ]
    add_table(doc, ['정책', '연계 내용'], policy_data, [5, 9])

    doc.add_paragraph()

    # ========== 6. 결론 ==========
    add_heading(doc, '6. 결론', level=1)

    add_paragraph(doc,
        "'온고지신 AI'는 검증된 AI 기술과 국내 유일의 대규모 치험례 데이터를 결합한 서비스임. "
        "MVP 개발이 완료되어 실제 동작하는 서비스를 보유하고 있으며, "
        "국내외 유사 서비스의 성공 사례를 통해 시장 가능성이 검증됨.")

    add_paragraph(doc,
        "정부의 한의학 디지털 전환 정책과 맞물려, 환자에게는 안전하고 효과적인 진료를, "
        "한의사에게는 진료 품질 향상을 위한 강력한 도구를 제공할 것임.")

    # 저장
    doc.save(OUTPUT_PATH)
    print(f"문서가 저장되었습니다: {OUTPUT_PATH}")
    return OUTPUT_PATH

if __name__ == "__main__":
    create_document()
