"""
실현 가능성(Solution) - 창업 아이템의 개발 계획 (핵심만)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = r"G:\내 드라이브\developer\hanmed-cdss"
SCREENSHOTS_DIR = os.path.join(BASE_DIR, "docs", "screenshots")
OUTPUT_PATH = os.path.join(BASE_DIR, "docs", "실현가능성_개발계획_final.docx")

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_paragraph(doc, text, bold=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_mixed_paragraph(doc, parts):
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return heading

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], 'E8F5E9')

    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    run.font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def add_image(doc, image_path, caption, width=5.0):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))

        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run(caption)
        caption_run.font.size = Pt(9)
        caption_run.font.name = '맑은 고딕'
        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        caption_run.italic = True
        caption_p.paragraph_format.space_after = Pt(10)

def create_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== 제목 ==========
    title = doc.add_heading('실현 가능성(Solution)_창업 아이템의 개발 계획', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== 1. 실현 가능성 ==========
    add_heading(doc, '1. 실현 가능성', level=1)

    # 1-1. 기술적 실현 가능성
    add_heading(doc, '1-1. 기술적 실현 가능성', level=2)

    add_mixed_paragraph(doc, [
        ("본 서비스는 ", False),
        ("이미 MVP(최소기능제품) 개발이 완료", True),
        ("되어 실제 동작하는 웹 서비스(https://hanmed-cdss.vercel.app)를 보유하고 있음.", False)
    ])

    add_image(doc, os.path.join(SCREENSHOTS_DIR, "03_dashboard.png"),
        "[그림 1] 온고지신 AI 대시보드 화면 (개발 완료)")

    tech_headers = ['항목', '현황', '근거']
    tech_data = [
        ['AI 기술', '상용화 단계', 'Claude API, GPT 등 LLM(대규모언어모델) 기술이 이미 성숙함'],
        ['데이터', '확보 완료', '6,000건 이상 치험례(치료경험기록) 디지털화 완료'],
        ['개발', 'MVP 완료', '웹 서비스가 실제 동작 중임'],
        ['인프라', '구축 완료', '클라우드 기반으로 안정적 운영 가능']
    ]
    add_table(doc, tech_headers, tech_data, [2.5, 3, 8.5])

    # 1-2. 시장 실현 가능성
    add_heading(doc, '1-2. 시장 실현 가능성 (롤모델 분석)', level=2)

    add_paragraph(doc, "AI 의료 서비스는 이미 국내외에서 성공적으로 운영되고 있어, 시장 실현 가능성이 검증됨.")

    rolemodel_headers = ['서비스명', '특징', '성과']
    rolemodel_data = [
        ['닥터앤서 2.0 (국내)', '정부 주도 AI 진단 보조 시스템', '12개 질환 대상, 전국 병원 도입'],
        ['루닛 (국내)', 'X-ray AI 분석', '글로벌 진출, 나스닥 상장'],
        ['Ping An Good Doctor (중국)', '전통의학 + AI 원격진료', '4억 명 사용자, 시가총액 10조원']
    ]
    add_table(doc, rolemodel_headers, rolemodel_data, [4, 5, 5])

    add_mixed_paragraph(doc, [
        ("특히 중국의 'Ping An Good Doctor'는 ", False),
        ("전통의학과 AI를 결합한 성공 사례", True),
        ("로, 본 사업의 방향성이 실현 가능함을 증명함.", False)
    ])

    doc.add_paragraph()

    # 1-3. 차별성
    add_heading(doc, '1-3. 롤모델 대비 차별성', level=2)

    diff_headers = ['항목', '기존 서비스', '온고지신 AI']
    diff_data = [
        ['대상 분야', '양방(서양의학) 중심', '한의학 전문 특화'],
        ['데이터 기반', '논문, 가이드라인', '실제 치험례 6,000건+ (국내 유일)'],
        ['처방 추천', '단순 정보 제공', '성공률 기반 AI 맞춤 추천'],
        ['상호작용 검사', '양약 간 검사만', '양약-한약 통합 검사']
    ]
    add_table(doc, diff_headers, diff_data, [3, 5.5, 5.5])

    add_mixed_paragraph(doc, [
        ("핵심 차별점: ", True),
        ("40년 경력 원로 한의사의 6,000건 이상 실제 치험례 데이터를 보유한 것은 국내에서 온고지신 AI가 유일함.", False)
    ])

    doc.add_paragraph()

    # ========== 2. 창업 아이템의 개발 계획 ==========
    add_heading(doc, '2. 창업 아이템의 개발 계획', level=1)

    # 2-1. 개발 로드맵
    add_heading(doc, '2-1. 단계별 개발 로드맵', level=2)

    roadmap_headers = ['단계', '시기', '주요 개발 내용', '완료 기준']
    roadmap_data = [
        ['Phase 1', '2026 Q1', 'MVP 고도화, 베타 테스트', '10개 한의원 시범 도입'],
        ['Phase 2', '2026 Q2', '정식 출시, EMR(전자의무기록) 연동', 'EMR 2개사 연동 완료'],
        ['Phase 3', '2026 하반기', '모바일 앱 출시, 기능 확장', '100개소 도입'],
        ['Phase 4', '2027', '서비스 고도화, 해외 진출 검토', '500개소 달성']
    ]
    add_table(doc, roadmap_headers, roadmap_data, [2, 2.5, 5.5, 4])

    # 2-2. 세부 개발 계획
    add_heading(doc, '2-2. 세부 개발 계획', level=2)

    add_paragraph(doc, "Phase 1: 핵심 기능 완성 (2026 Q1)", bold=True)
    phase1 = [
        "• AI 변증 분석 정확도 향상 (목표: 90% 이상)",
        "• 학파별(고방, 후세방, 사상방) 처방 분류 시스템 구축",
        "• 베타 테스터 피드백 수집 및 UI/UX 개선"
    ]
    for item in phase1:
        add_paragraph(doc, item, space_after=3)

    add_image(doc, os.path.join(SCREENSHOTS_DIR, "05_consultation.png"),
        "[그림 2] AI 진료 어시스턴트 화면 (Phase 1 고도화 대상)")

    add_paragraph(doc, "Phase 2: 시장 확장 (2026 Q2)", bold=True)
    phase2 = [
        "• 주요 EMR 업체(한의타임, 오아시스 등)와 연동 API 개발",
        "• 구독 결제 시스템 구축 및 정식 서비스 출시",
        "• 고객 지원 체계(CS) 구축"
    ]
    for item in phase2:
        add_paragraph(doc, item, space_after=3)

    add_paragraph(doc, "Phase 3: 서비스 확장 (2026 하반기)", bold=True)
    phase3 = [
        "• iOS/Android 모바일 앱 개발",
        "• 음성 입력 기능 추가 (진료 중 핸즈프리 사용)",
        "• 팔강변증(8가지 기준 진단법) 자동 분석기 개발"
    ]
    for item in phase3:
        add_paragraph(doc, item, space_after=3)

    add_image(doc, os.path.join(SCREENSHOTS_DIR, "09_interactions.png"),
        "[그림 3] 양약-한약 상호작용 검사 화면 (Phase 3 고도화 대상)")

    add_paragraph(doc, "Phase 4: 고도화 및 확장 (2027)", bold=True)
    phase4 = [
        "• AI 모델 자체 학습 기능 개발 (사용자 피드백 기반)",
        "• 해외 시장 진출을 위한 다국어 지원",
        "• 대형 한방병원 및 한의과대학 대상 엔터프라이즈 버전 개발"
    ]
    for item in phase4:
        add_paragraph(doc, item, space_after=3)

    doc.add_paragraph()

    # 2-3. 기술 스택
    add_heading(doc, '2-3. 기술 스택', level=2)

    tech_stack_headers = ['구분', '기술', '선정 이유']
    tech_stack_data = [
        ['AI 엔진', 'Claude API + RAG', '자연어 이해 능력 우수, 한의학 맥락 파악에 적합'],
        ['프론트엔드', 'React + TypeScript', '안정적인 웹 개발, 유지보수 용이'],
        ['백엔드', 'NestJS + PostgreSQL', '확장성 높은 서버 구조, 대용량 데이터 처리'],
        ['배포', 'Vercel + Fly.io', '클라우드 기반 자동 확장, 99.9% 가동률']
    ]
    add_table(doc, tech_stack_headers, tech_stack_data, [3, 4, 7])

    # 저장
    doc.save(OUTPUT_PATH)
    print(f"문서가 저장되었습니다: {OUTPUT_PATH}")
    return OUTPUT_PATH

if __name__ == "__main__":
    create_document()
