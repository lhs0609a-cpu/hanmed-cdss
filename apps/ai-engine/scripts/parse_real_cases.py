"""
실제 치험례 파싱 스크립트
치험례/word 폴더의 docx 파일들에서 치험례를 추출하여 JSON으로 저장
"""

import json
import re
import os
from pathlib import Path
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, asdict, field
from docx import Document
from datetime import datetime
import uuid

# 프로젝트 경로
PROJECT_ROOT = Path(__file__).parent.parent.parent.parent
CASES_DIR = PROJECT_ROOT / "치험례" / "word"
OUTPUT_DIR = PROJECT_ROOT / "apps" / "ai-engine" / "data"


@dataclass
class RealCase:
    """실제 치험례 데이터"""
    id: str
    source_file: str
    formula_name: str = ""
    formula_hanja: str = ""
    case_code: str = ""  # 42-01-01 등
    title: str = ""
    chief_complaint: str = ""
    symptoms: List[str] = field(default_factory=list)
    sub_symptoms: List[str] = field(default_factory=list)
    diagnosis: str = ""
    patient_name: str = ""
    patient_age: Optional[int] = None
    patient_gender: Optional[str] = None
    patient_constitution: Optional[str] = None  # 소음인/태음인/소양인/태양인
    patient_occupation: str = ""
    patient_address: str = ""
    appearance: str = ""  # 용모
    history: str = ""  # 과정/병력
    reference: str = ""  # 참고
    differentiation: str = ""  # 변상/변증
    treatment_principle: str = ""  # 치법
    prescription_plan: str = ""  # 처방구상
    medications: List[Dict] = field(default_factory=list)  # 투약 기록들
    progress: List[Dict] = field(default_factory=list)  # 경과 기록들
    result: str = ""  # 최종 결과
    full_text: str = ""  # 전체 원문
    data_source: str = "real_clinical_case"
    created_at: str = ""
    search_text: str = ""


def extract_age(text: str) -> Optional[int]:
    """텍스트에서 나이 추출"""
    patterns = [
        r'(\d+)\s*세',
        r'(\d+)\s*歲',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            age = int(match.group(1))
            if 0 < age < 120:
                return age
    return None


def extract_gender(text: str) -> Optional[str]:
    """텍스트에서 성별 추출"""
    if '여' in text or '女' in text:
        return 'F'
    if '남' in text or '男' in text:
        return 'M'
    return None


def extract_constitution(text: str) -> Optional[str]:
    """텍스트에서 체질 추출"""
    constitutions = ['소음인', '태음인', '소양인', '태양인']
    for const in constitutions:
        if const in text:
            return const
    return None


def extract_formula_name(text: str) -> Optional[str]:
    """텍스트에서 처방명 추출"""
    # 제외할 단어 (처방명이 아닌 것)
    exclude_words = ['처방', '투약', '복용', '증상', '경과', '참고', '변증', '치법',
                     '가감', '합방', '용량', '약제', '주증', '부증', '원방', '본방',
                     '이라고', '라고', '것이라고', '한다고']

    # 잘못된 패턴 (한국어 어미로 끝나는 것들) - "~고" 처리
    false_endings = ['하고', '다고', '라고', '아고', '하고', '이고', '으로', '없고',
                     '있고', '했고', '됐고', '못하고', '않고', '치고', '되고', '지고',
                     '여위고', '아프고', '것이고', '복용하고', '수술하고', '먹고',
                     '순환', '한의원', '의원', '약국', '약방', '병원', '침술']

    def is_valid_formula(name: str) -> bool:
        """유효한 처방명인지 확인"""
        if not name or len(name) < 3:
            return False
        if name in exclude_words:
            return False
        # 일반 동사/형용사 어미로 끝나는지 확인
        for ending in false_endings:
            if name.endswith(ending):
                return False
        # "고"로 끝나는 경우 특별 검증 (경옥고, 자옥고 등만 허용)
        if name.endswith('고') and len(name) >= 3:
            # 유효한 "~고" 처방명들
            valid_go = ['경옥고', '자옥고', '응약고', '황련고', '자운고', '옥용고']
            if name not in valid_go:
                return False
        return True

    # 한약 처방명 패턴 (탕, 산, 환, 단, 음, 원, 전 등으로 끝남 - "고" 제외)
    formula_patterns = [
        # 복용/처방 등 뒤에 오는 처방명
        r'(?:복용|투약|처방|사용)\s*[:：]?\s*([가-힣]{2,10}(?:탕|산|환|단|음|원|전|방|제))',
        # 일반적인 처방명 패턴
        r'([가-힣]{2,10}(?:탕|산|환|단|음|원|전|방|제))\s*(?:을|를|가|이|의|으로|에|처방)?',
        r'([가-힣]{2,10}(?:탕|산|환|단|음|원|전|방|제))\s*\d*\s*(?:첩|일분|제)?',
        # 가감/합방 패턴
        r'([가-힣]{2,10}(?:탕|산|환|단|음|원|전|방|제))\s*(?:가감|합방)',
    ]

    # 단일 약재/제제 패턴 (특수 제제들)
    single_patterns = [
        r'(가시오가피(?:액|탕액)?)',
        r'(녹용대보탕)',
        r'(우황청심원)',
        r'(공진단)',
        r'(경옥고)',
        r'(자옥고)',
        r'(쌍화탕)',
        r'(십전대보탕)',
        r'(보중익기탕)',
        r'(팔물탕)',
        r'(사물탕)',
        r'(사군자탕)',
        r'(소시호탕)',
        r'(반하사심탕)',
        r'(육군자탕)',
        r'(귀비탕)',
        r'(청심연자탕)',
    ]

    # 주요 처방명 패턴 매칭
    for pattern in formula_patterns:
        matches = re.findall(pattern, text)
        if matches:
            for match in matches:
                if is_valid_formula(match):
                    return match

    # 단일 패턴 매칭
    for pattern in single_patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1)

    # 마지막 시도: 더 넓은 패턴 (고 제외)
    broad_pattern = r'([가-힣]{3,12}(?:탕|산|환|단|음|원|전|방|제))'
    matches = re.findall(broad_pattern, text)
    for match in matches:
        if is_valid_formula(match) and len(match) >= 4:
            return match

    return None


def extract_all_formulas(text: str) -> List[str]:
    """텍스트에서 모든 처방명 추출"""
    formulas = set()

    # 한약 처방명 패턴
    pattern = r'([가-힣]{2,10}(?:탕|산|환|고|단|음|원|전|방|제))'
    matches = re.findall(pattern, text)

    # 필터링 (일반 단어 제외)
    exclude_words = ['처방', '투약', '복용', '증상', '경과', '참고', '변증', '치법']
    for m in matches:
        if m not in exclude_words and len(m) >= 3:
            formulas.add(m)

    # 단일 약재/제제
    single_formulas = ['가시오가피', '녹용', '우황청심원', '공진단', '경옥고', '양배추']
    for sf in single_formulas:
        if sf in text:
            formulas.add(sf)

    return list(formulas)


def parse_sasang_case_file(filepath: Path) -> List[RealCase]:
    """사상체질 치험례 모음 파일 파싱"""
    cases = []
    doc = Document(filepath)

    full_text = '\n'.join([p.text for p in doc.paragraphs])

    # ■ 로 시작하는 케이스들 분리
    case_pattern = r'■\s*([^■]+)'
    case_blocks = re.findall(case_pattern, full_text, re.DOTALL)

    for block in case_blocks:
        if len(block.strip()) < 50:
            continue

        case = RealCase(
            id=str(uuid.uuid4())[:8],
            source_file=filepath.name,
            created_at=datetime.now().isoformat(),
        )

        # 첫 줄에서 처방명과 주증 추출
        first_line = block.strip().split('\n')[0]

        # 형방지황탕(42-01-01) -- 후두염(候頭炎) 불면(不眠)
        # 처방명 패턴: 탕/산/환/단/음/원/전/방으로 끝나야 함
        formula_match = re.match(r'([가-힣]+(?:탕|산|환|단|음|원|전|방))\s*\(?([\d\-]+)?\)?', first_line)
        if formula_match:
            case.formula_name = formula_match.group(1)
            if formula_match.group(2):
                case.case_code = formula_match.group(2)

        # 주증 추출
        if '--' in first_line:
            case.title = first_line.split('--')[-1].strip()
            case.chief_complaint = case.title

        # 환자 정보 추출 (두번째 줄 근처)
        patient_pattern = r'([가-힣\s○0]+)\s+(여|남)\s*(\d+)세?\s+(소음인|태음인|소양인|태양인)?'
        patient_match = re.search(patient_pattern, block)
        if patient_match:
            case.patient_name = patient_match.group(1).strip()
            case.patient_gender = 'F' if patient_match.group(2) == '여' else 'M'
            case.patient_age = int(patient_match.group(3))
            if patient_match.group(4):
                case.patient_constitution = patient_match.group(4)

        # 체질 추출 (없으면 파일명에서)
        if not case.patient_constitution:
            case.patient_constitution = extract_constitution(block)
            if not case.patient_constitution:
                if '소양인' in filepath.name:
                    case.patient_constitution = '소양인'
                elif '소음인' in filepath.name:
                    case.patient_constitution = '소음인'
                elif '태음인' in filepath.name:
                    case.patient_constitution = '태음인'
                elif '태양인' in filepath.name:
                    case.patient_constitution = '태양인'

        # 섹션별 추출
        sections = {
            '용모': ('appearance', r'용\s*모\s*[:：]?\s*(.+?)(?=￭|$)'),
            '과정': ('history', r'과\s*정\s*[:：]?\s*(.+?)(?=￭|$)'),
            '주증상': ('chief_complaint', r'주\s*증\s*상\s*[:：]?\s*(.+?)(?=￭|부수증상|참고|$)'),
            '부수증상': ('sub_symptoms_text', r'부수증상\s*[:：]?\s*(.+?)(?=￭|참고|$)'),
            '참고': ('reference', r'참\s*고\s*[:：]?\s*(.+?)(?=￭|변상|변증|$)'),
            '변상': ('differentiation', r'변\s*상\s*[:：]?\s*(.+?)(?=￭|변증|치법|$)'),
            '변증': ('differentiation', r'변\s*증\s*\d*\s*[:：]?\s*(.+?)(?=￭|치법|처방|$)'),
            '치법': ('treatment_principle', r'치\s*법\s*[:：]?\s*(.+?)(?=￭|처방|투약|$)'),
            '처방구상': ('prescription_plan', r'처방구상\s*\d*\s*[:：]?\s*(.+?)(?=￭|투약|$)'),
        }

        for section_name, (field_name, pattern) in sections.items():
            match = re.search(pattern, block, re.DOTALL)
            if match:
                value = match.group(1).strip()
                if field_name == 'sub_symptoms_text':
                    # 증상 리스트로 분리
                    symptoms = re.findall(r'\d+\.\s*(.+?)(?=\d+\.|$)', value)
                    case.sub_symptoms = [s.strip() for s in symptoms if s.strip()]
                else:
                    setattr(case, field_name, value[:500])  # 길이 제한

        # 투약/경과 추출
        medication_pattern = r'투\s*약\s*(\d*)\s*[:：]?\s*(.+?)(?=￭|경과|$)'
        progress_pattern = r'경\s*과\s*(\d*)\s*[:：]?\s*(.+?)(?=￭|투약|검토|$)'

        for match in re.finditer(medication_pattern, block, re.DOTALL):
            case.medications.append({
                'order': match.group(1) or '1',
                'content': match.group(2).strip()[:300]
            })

        for match in re.finditer(progress_pattern, block, re.DOTALL):
            case.progress.append({
                'order': match.group(1) or '1',
                'content': match.group(2).strip()[:300]
            })

        # 결과 추출 (마지막 경과에서)
        if case.progress:
            case.result = case.progress[-1].get('content', '')[:200]

        # 전체 텍스트
        case.full_text = block[:2000]

        # 처방명이 없으면 전체 텍스트에서 추출 시도
        if not case.formula_name:
            case.formula_name = extract_formula_name(block) or ""

        # 검색 텍스트 생성
        case.search_text = ' '.join([
            case.formula_name,
            case.title,
            case.chief_complaint,
            ' '.join(case.symptoms),
            ' '.join(case.sub_symptoms),
            case.diagnosis,
            case.patient_constitution or '',
        ])

        if case.formula_name or case.chief_complaint:
            cases.append(case)

    return cases


def parse_individual_case_file(filepath: Path) -> List[RealCase]:
    """개별 치험례 파일 파싱 (00-xxx.docx 등)"""
    cases = []
    doc = Document(filepath)

    full_text = '\n'.join([p.text for p in doc.paragraphs])

    case = RealCase(
        id=str(uuid.uuid4())[:8],
        source_file=filepath.name,
        created_at=datetime.now().isoformat(),
    )

    # 환자 정보 추출 (첫 줄들에서)
    lines = full_text.split('\n')
    for line in lines[:10]:
        # 오 태 환   남 72   소양인
        patient_match = re.search(r'([가-힣\s]+)\s+(남|여)\s*(\d+)\s+(소음인|태음인|소양인|태양인)?', line)
        if patient_match:
            case.patient_name = patient_match.group(1).strip()
            case.patient_gender = 'M' if patient_match.group(2) == '남' else 'F'
            case.patient_age = int(patient_match.group(3))
            if patient_match.group(4):
                case.patient_constitution = patient_match.group(4)
            break

    if not case.patient_constitution:
        case.patient_constitution = extract_constitution(full_text)

    # 주증상 추출
    symptoms_match = re.search(r'주\s*증\s*상\s*[:：]?\s*(.+?)(?=참\s*고|투약|복용|$)', full_text, re.DOTALL)
    if symptoms_match:
        case.chief_complaint = symptoms_match.group(1).strip()[:300]
        # 증상 리스트 추출
        symptom_items = re.findall(r'\d+\.\s*(.+?)(?=\d+\.|$)', case.chief_complaint)
        case.symptoms = [s.strip() for s in symptom_items if s.strip()]

    # 참고 추출
    ref_match = re.search(r'참\s*고\s*[:：]?\s*(.+?)(?=투약|복용|$)', full_text, re.DOTALL)
    if ref_match:
        case.reference = ref_match.group(1).strip()[:500]

    # 복용/경과 추출
    medication_pattern = r'복\s*용\s*(\d*)\s*[:：]?\s*(.+?)(?=경\s*과|복\s*용|$)'
    progress_pattern = r'경\s*과\s*(\d*)\s*[:：]?\s*(.+?)(?=복\s*용|증\s*상|$)'

    for match in re.finditer(medication_pattern, full_text, re.DOTALL):
        case.medications.append({
            'order': match.group(1) or '1',
            'content': match.group(2).strip()[:300]
        })

    for match in re.finditer(progress_pattern, full_text, re.DOTALL):
        case.progress.append({
            'order': match.group(1) or '1',
            'content': match.group(2).strip()[:300]
        })

    # 결과 추출
    if case.progress:
        case.result = case.progress[-1].get('content', '')[:200]

    # 처방명 추출 (전체 텍스트에서)
    case.formula_name = extract_formula_name(full_text) or ""

    # 타이틀 설정
    case.title = case.chief_complaint[:50] if case.chief_complaint else filepath.stem

    # 전체 텍스트
    case.full_text = full_text[:2000]

    # 검색 텍스트
    case.search_text = ' '.join([
        case.formula_name,
        case.title,
        case.chief_complaint,
        ' '.join(case.symptoms),
        case.patient_constitution or '',
    ])

    if case.chief_complaint or case.symptoms:
        cases.append(case)

    return cases


def parse_collection_file(filepath: Path) -> List[RealCase]:
    """모음집 파일 파싱 (고령자채록, 감초당 등)"""
    cases = []
    doc = Document(filepath)

    full_text = '\n'.join([p.text for p in doc.paragraphs])

    # 케이스 구분자 찾기 (번호나 특수문자로 구분)
    # 다양한 패턴 시도
    patterns = [
        r'(?:^|\n)(\d+\.\s*.+?)(?=\n\d+\.|$)',  # 1. 케이스
        r'(?:^|\n)(【.+?】.+?)(?=【|$)',  # 【케이스】
        r'(?:^|\n)(▣.+?)(?=▣|$)',  # ▣ 케이스
    ]

    case_blocks = []
    for pattern in patterns:
        blocks = re.findall(pattern, full_text, re.DOTALL)
        if len(blocks) > 3:  # 최소 3개 이상 찾았으면 사용
            case_blocks = blocks
            break

    # 패턴으로 못 찾으면 문단 단위로 분리
    if not case_blocks:
        # 환자 정보 패턴으로 분리
        case_blocks = re.split(r'(?=\n[가-힣\s]+\s+(?:남|여)\s+\d+세?)', full_text)

    for block in case_blocks:
        if len(block.strip()) < 100:
            continue

        case = RealCase(
            id=str(uuid.uuid4())[:8],
            source_file=filepath.name,
            created_at=datetime.now().isoformat(),
        )

        # 환자 정보 추출
        patient_match = re.search(r'([가-힣○\s]+)\s+(남|여)\s*(\d+)세?\s*(소음인|태음인|소양인|태양인)?', block)
        if patient_match:
            case.patient_name = patient_match.group(1).strip()
            case.patient_gender = 'M' if patient_match.group(2) == '남' else 'F'
            case.patient_age = int(patient_match.group(3))
            if patient_match.group(4):
                case.patient_constitution = patient_match.group(4)

        if not case.patient_constitution:
            case.patient_constitution = extract_constitution(block)

        # 주증상
        symptoms_match = re.search(r'(?:주\s*증|증\s*상|호\s*소)\s*[:：]?\s*(.+?)(?=처방|치료|투약|$)', block, re.DOTALL)
        if symptoms_match:
            case.chief_complaint = symptoms_match.group(1).strip()[:300]

        # 처방 - "고"는 유효한 처방명(경옥고 등)만 허용
        formula_match = re.search(r'처\s*방\s*[:：]?\s*([가-힣]+(?:탕|산|환|단|음|원|전|방|제))', block)
        if formula_match:
            case.formula_name = formula_match.group(1)
        else:
            # 패턴 매칭 실패시 전체 텍스트에서 추출
            case.formula_name = extract_formula_name(block) or ""

        # 결과
        result_match = re.search(r'(?:경과|결과|효과)\s*[:：]?\s*(.+?)(?=\n\n|$)', block, re.DOTALL)
        if result_match:
            case.result = result_match.group(1).strip()[:200]

        case.title = case.chief_complaint[:50] if case.chief_complaint else f"케이스-{len(cases)+1}"
        case.full_text = block[:2000]
        case.search_text = ' '.join([
            case.formula_name,
            case.chief_complaint,
            case.patient_constitution or '',
        ])

        if case.chief_complaint or case.formula_name:
            cases.append(case)

    return cases


def main():
    """메인 실행"""
    print("=" * 60)
    print("실제 치험례 파싱 스크립트")
    print("=" * 60)

    if not CASES_DIR.exists():
        print(f"치험례 폴더가 없습니다: {CASES_DIR}")
        return

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    all_cases = []

    # 파일별 처리
    docx_files = list(CASES_DIR.glob("*.docx"))
    print(f"총 {len(docx_files)}개 docx 파일 발견")

    for filepath in docx_files:
        print(f"\n처리 중: {filepath.name}")

        try:
            # 파일 유형에 따라 다른 파서 사용
            if '사상' in filepath.name and '치험례' in filepath.name:
                cases = parse_sasang_case_file(filepath)
                print(f"  → 사상체질 치험례: {len(cases)}개 추출")
            elif re.match(r'^\d', filepath.name) or filepath.name.startswith('00-'):
                cases = parse_individual_case_file(filepath)
                print(f"  → 개별 치험례: {len(cases)}개 추출")
            elif '채록' in filepath.name or '모음' in filepath.name:
                cases = parse_collection_file(filepath)
                print(f"  → 모음집 치험례: {len(cases)}개 추출")
            elif '태' in filepath.name and '장' in filepath.name:
                cases = parse_sasang_case_file(filepath)
                print(f"  → 태음인 치험례: {len(cases)}개 추출")
            else:
                cases = parse_collection_file(filepath)
                print(f"  → 기타 치험례: {len(cases)}개 추출")

            all_cases.extend(cases)

        except Exception as e:
            print(f"  → 오류: {e}")

    print(f"\n총 {len(all_cases)}개 치험례 추출 완료")

    # 통계
    with_constitution = sum(1 for c in all_cases if c.patient_constitution)
    with_formula = sum(1 for c in all_cases if c.formula_name)
    with_result = sum(1 for c in all_cases if c.result)

    print(f"\n통계:")
    print(f"  - 체질 정보: {with_constitution}개 ({with_constitution/len(all_cases)*100:.1f}%)")
    print(f"  - 처방 정보: {with_formula}개 ({with_formula/len(all_cases)*100:.1f}%)")
    print(f"  - 결과 정보: {with_result}개 ({with_result/len(all_cases)*100:.1f}%)")

    # JSON 저장
    output_file = OUTPUT_DIR / "real_clinical_cases.json"
    cases_data = [asdict(c) for c in all_cases]

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(cases_data, f, ensure_ascii=False, indent=2)

    print(f"\n저장 완료: {output_file}")

    # 샘플 출력
    if all_cases:
        print(f"\n샘플 (첫 번째):")
        sample = all_cases[0]
        print(f"  제목: {sample.title}")
        print(f"  처방: {sample.formula_name}")
        print(f"  환자: {sample.patient_age}세 {sample.patient_gender} {sample.patient_constitution or ''}")
        print(f"  주소증: {sample.chief_complaint[:100]}")
        print(f"  결과: {sample.result[:100]}")


if __name__ == "__main__":
    main()
